from flask import Flask, render_template, request, send_file, redirect, url_for, flash, make_response, jsonify, session
import pandas as pd
import os
import subprocess
import time
import pyautogui
from functools import wraps
from werkzeug.utils import secure_filename
import win32com.client
from docx import Document
from datetime import datetime, timedelta
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from webdriver_manager.chrome import ChromeDriverManager
import threading
import logging
import re
import json
import requests

app = Flask(__name__)
app.secret_key = 'your_secret_key_change_this_in_production'
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['MAX_CONTENT_LENGTH'] = 25 * 1024 * 1024  # 25MB max file size

# Hardcoded download filename and path
app.config['DOWNLOAD_FILENAME'] = 'ROB.xlsx'
app.config['DOWNLOAD_PATH'] = r'C:\Users\akshat\Desktop\RPA\\' + app.config['DOWNLOAD_FILENAME']

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Set up logging to capture output
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Add a global variable to track processing status
processing_status = {
    'active': False,
    'message': 'Ready',
    'progress': 0,
    'total': 0,
    'current_file': '',
    'logs': []
}

def allowed_file(filename):
    """Check if file extension is allowed"""
    ALLOWED_EXTENSIONS = {'xlsx', 'xls', 'csv'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_backend_file(filename):
    """Check if backend file extension is allowed"""
    BACKEND_EXTENSIONS = {'xlsx', 'xls'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in BACKEND_EXTENSIONS

def log_to_status(message):
    """Add a message to the processing status logs"""
    global processing_status
    processing_status['logs'].append(f"{datetime.now().strftime('%H:%M:%S')}: {message}")
    print(f"[LOG] {message}")

# ============================================================================
# HOME ROUTE
# ============================================================================

@app.route('/')
def home():
    return render_template('home.html')
@app.route('/index.html')
def index():
    return render_template('index.html')

# ============================================================================
# DOCUMENT PROCESSING ROUTES
# ============================================================================

@app.route('/document_processing', methods=['GET', 'POST'])
def document_processing():
    global processing_status
    
    if request.method == 'POST':
        try:
            # Get form data - use session data as defaults if available
            article_code = request.form.get('article_code') or request.form.get('open_pr_id') or session.get('open_pr_id', '6HA-2025-M6K439')
            author_name = request.form.get('author_name') or session.get('username', 'akshat tiwari')
            author_email = request.form.get('author_email') or session.get('email', 'akshat@coherentmarketinsights.com')
            company_name = request.form.get('company_name', 'Coherent Market Insights')
            phone_number = request.form.get('phone_number') or session.get('mobile', '1234567890')
            image_path = request.form.get('image_path')  # Get image path from form
            
            # Power Automate output folder path
            custom_folder = request.form.get('custom_folder')
            if custom_folder:
                folder_path = custom_folder
            else:
                today = datetime.today()
                folder_path = rf'C:\Users\akshat\Desktop\RPA\Files\{today.year}\{today.strftime("%m")}\{today.strftime("%d")}'
            
            processing_mode = request.form.get('processing_mode', 'manual')
            
            # Validate paths before processing
            excel_path = r'C:\Users\akshat\Desktop\RPA\ROB.xlsx'
            
            # Check if required files exist
            validation_errors = []
            if not os.path.exists(excel_path):
                validation_errors.append(f"Excel file not found: {excel_path}")
            if not os.path.exists(folder_path):
                validation_errors.append(f"Folder not found: {folder_path}")
            if image_path and not os.path.exists(image_path):
                validation_errors.append(f"Image file not found: {image_path}")
            if not image_path:
                validation_errors.append("Image path is required")
            
            if validation_errors:
                for error in validation_errors:
                    flash(error)
                return render_template('document_processing.html', 
                                     session_data={
                                         'username': session.get('username', ''),
                                         'email': session.get('email', ''),
                                         'mobile': session.get('mobile', ''),
                                         'open_pr_id': session.get('open_pr_id', ''),
                                         'image_path': image_path or ''
                                     })
            
            # Reset processing status
            processing_status = {
                'active': True,
                'message': 'Starting processing...',
                'progress': 0,
                'total': 0,
                'current_file': '',
                'logs': []
            }
            
            # Start processing in background thread - NOW INCLUDING image_path
            if processing_mode == 'auto':
                threading.Thread(target=process_documents_auto_with_feedback, 
                               args=(folder_path, article_code, author_name, author_email, 
                                    company_name, phone_number, image_path)).start()
            else:
                threading.Thread(target=process_documents_manual_with_feedback, 
                               args=(folder_path, article_code, author_name, author_email, 
                                    company_name, phone_number, image_path)).start()
            
            flash('Processing started! Check the status page for updates.')
            return redirect(url_for('processing_status'))
            
        except Exception as e:
            flash(f'Error starting processing: {str(e)}')
            logger.error(f"Error in document_processing: {e}")
            return render_template('document_processing.html', 
                                 session_data={
                                     'username': session.get('username', ''),
                                     'email': session.get('email', ''),
                                     'mobile': session.get('mobile', ''),
                                     'open_pr_id': session.get('open_pr_id', ''),
                                     'image_path': request.form.get('image_path', '')
                                 })
    
    # Pre-populate form with session data if available
    return render_template('document_processing.html', 
                         session_data={
                             'username': session.get('username', ''),
                             'email': session.get('email', ''),
                             'mobile': session.get('mobile', ''),
                             'open_pr_id': session.get('open_pr_id', ''),
                             'image_path': session.get('image_path', '')
                         })

@app.route('/processing_status')
def processing_status_page():
    return render_template('processing_status.html')

@app.route('/api/get_processing_status')
def get_processing_status():
    """API endpoint to get current processing status"""
    global processing_status
    return jsonify(processing_status)

# ============================================================================
# DOCUMENT PROCESSING FUNCTIONS
# ============================================================================
import win32com.client
import re

import win32com.client
import re

def text_of_press_release(doc_path):
    # Open Word application
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # Run in the background

    # Open the .doc file (adjust the file path if necessary)
    doc2 = word.Documents.Open(doc_path)

    # Extract the entire text from the document
    doc_text = doc2.Content.Text

    # Remove the first line from the document text
    lines = doc_text.splitlines()
    if len(lines) > 1:
        doc_text = '\n'.join(lines[1:])

    # Define the headings for which you want to add line breaks
    headings = [
        "➤Market Size and Overview",
        "➤Actionable Insights",
        "➤Actionable insights",
        "➤Growth factors",
        "➤Growth Factors",
        "➤Market trends",
        "➤Market Trends",
        "➤Key takeaways ",
        "➤Key Takeaways",
        "➤Market Segment and Regional Coverage ",
        "➤Market segment and regional coverage",
        "➤Key players",
        "➤Key Players",
        "➤Competitive Strategies and Outcomes",
        "❓ Frequently Asked Questions",
        "❓ Frequently asked questions"
    ]

    # FIXED: Add a line space AFTER each heading (not before and after)
    for heading in headings:
        doc_text = doc_text.replace(heading, f"{heading}\n")

    # Define the regex pattern for URLs
    url_pattern = re.compile(r"(https?://[^\s]+)")
    
    # Define regex patterns for FAQ questions (numbered questions and roman numerals)
    faq_pattern_numbers = re.compile(r"^\d+\.\s")  # Matches "1. ", "2. ", etc.
    faq_pattern_roman = re.compile(r"^[ivxlcdmIVXLCDM]+\.\s")  # Matches "i. ", "ii. ", "I. ", "II. ", etc.
    
    # Define regex pattern for CTA links (➔)
    cta_pattern = re.compile(r"^➔")  # Matches lines starting with ➔

    # Split the text into lines
    lines = doc_text.splitlines()
    processed_lines = []

    # Iterate over each line
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        
        # Skip empty lines in processing, we'll add them strategically
        if not line_stripped:
            continue
            
        # Check if this is a CTA line
        is_cta = cta_pattern.match(line_stripped)
        
        # Check if previous line was a CTA (for adjacent CTA handling)
        prev_was_cta = False
        if processed_lines:
            last_non_empty = None
            for prev_line in reversed(processed_lines):
                if prev_line.strip():
                    last_non_empty = prev_line.strip()
                    break
            if last_non_empty and cta_pattern.match(last_non_empty):
                prev_was_cta = True
        
        # Check if this line is a heading (starts with ➤ or ❓)
        is_heading = line_stripped.startswith('➤') or line_stripped.startswith('❓')
        
        # If a line contains a URL, add space before and after the URL
        if url_pattern.search(line):
            # Add space before (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            processed_lines.append('')  # Add space after
            
        # If a line is an FAQ question (starts with number or roman numeral), add space before it
        elif faq_pattern_numbers.match(line_stripped) or faq_pattern_roman.match(line_stripped):
            # Add space before FAQ question (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            
        # If this is a CTA line
        elif is_cta:
            # Add space before CTA (unless previous was also CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            
        # If this line is a heading (starts with ➤ or ❓)
        elif is_heading:
            # Add space before heading (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            # FIXED: Add space AFTER heading
            processed_lines.append('')
            
        else:
            # Regular content line
            processed_lines.append(line)

    # Join the processed lines back into a single string
    chunk = "\n".join(processed_lines)
    
    # Clean up multiple consecutive empty lines (replace with single empty line)
    chunk = re.sub(r'\n\s*\n\s*\n+', '\n\n', chunk)

    # Close the document
    doc2.Close()
    word.Quit()

    # Return the processed content
    return chunk
def run_selenium_automation_single(row_data, category, article_code, author_name, author_email, company_name, phone_number, image_path):
    """Run Selenium automation for a single press release submission"""
    try:
        import random
        AUTHOR_DESCRIPTIONS = [
    """Author of this marketing PR:
Ravina Pandya, Content Writer, has a strong foothold in the market research industry. She specializes in writing well-researched articles from different industries, including food and beverages, information and technology, healthcare, chemical and materials, etc.
 """    
    """ Author of this marketing PR :
Money Singh is a seasoned content writer with over four years of experience in the market research sector. Her expertise spans various industries, including food and beverages, biotechnology, chemical and materials, defense and aerospace, consumer goods, etc. 
"""    
    """ Author of this marketing PR:

Alice Mutum is a seasoned senior content editor at Coherent Market Insights, leveraging extensive expertise gained from Openpr her previous role as a content writer. With seven years in content development, Alice masterfully employs SEO best practices and cutting-edge digital marketing strategies to craft high-ranking, impactful content. As an editor, she meticulously ensures flawless grammar and punctuation, precise data accuracy, and perfect alignment with audience needs in every research report. Alice's dedication to excellence and her strategic approach to content make her an invaluable asset in the world of market insights.
"""]
        # Extract data from the row
        market_name = row_data['Market Name']
        category = row_data['Category']
        TITLE_PROMPTS = [
    "Is Booming Worldwide 2025-2032",
    "Generated Opportunities, Future Scope 2025-2032",
    "Future Business Opportunities 2025-2032",
    "Growth in Future Scope 2025-2032",
    "Is Booming So Rapidly Growth by 2032",
    "Is Booming So Rapidly with CAGR of 6.9%",
    "An Analysis of Size, Shares, Business Growth, and Upcoming Trends Forecast 2025-2032",
    "2025-2032 Emerging Trends in Industry Dynamics, Size Insights, Share, and Future Growth",
    "Set to Witness Significant Growth by 2025-2032",
    "to Witness Massive Growth by 2032",
    "Size, Share 2025 Analysis of Rising Business Opportunities with Prominent Investment, Forecast to 2032",
    "by Trends, Dynamic Innovation in Technology and 2032 Forecast, Opportunities, and Challenges, Trends",
    "Size 2025 Emerging Demands, Share, Trends, Futuristic Opportunity, Share and Forecast To 2032",
    "Size, Share Growth Status, Emerging Technology, Key Players, Industry Challenges, and Forecast till 2032",
    "Insights 2025-2032: Global Expansion, Revenue Trends, and Strategic Growth Plans",
    "Outlook 2025-2032: Growth Drivers, Share, And Trends As Revealed In New Report"
]

        
        # Extract companies covered (and ensure it handles any missing or malformed data)
        companies = row_data.get('Companies covered', 'No companies listed')
        log_to_status(f"Trying to '{companies}' for market '{market_name}'")

        # Create article title from market name and companies
        # If companies are covered, limit to the first 5 companies, otherwise just use market name
        if companies and isinstance(companies, str) and companies.strip():
            company_list = [c.strip() for c in companies.split(',') if c.strip()]
            first_five_companies = ', '.join(company_list[:5]) if company_list else 'No companies available'
        else:
            first_five_companies = 'No companies available'
        random_prompt = random.choice(TITLE_PROMPTS)
        article_title = f"{market_name} {random_prompt} {first_five_companies}"
        
        # Create multiline text from the row data
        multiline_text= f"""
{market_name} - Market Insights Report

Market Overview:
{row_data.get('Market Size', 'Market analysis and insights')}

Forecast Period: {row_data.get('Forecast Period', 'N/A')}
CAGR: {row_data.get('CAGR', 'N/A')}

Key Market Players:
{row_data.get('Key Players', 'Leading companies in the market')}

For more detailed information, please refer to our comprehensive market research report.
        """
        
        log_to_status(f"Processing: {market_name}")
        log_to_status(f"Using category: {category}")

        log_to_status("Starting Selenium automation for: " + market_name)
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        options.add_argument("--start-maximized")
        
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.openpr.com/')
        
        # Handle cookie consent
        try:
            reject = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="cmpbntnotxt"]'))
            )
            reject.click()
        except:
            pass
        
        # Navigate to submit page
        submit = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="navbarText"]/ul/li[3]/a'))
        )
        submit.click()
        
        # Enter article code
        input_box = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="code"]'))
        )
        input_box.clear()
        input_box.send_keys(article_code)
        
        # Submit code
        try:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(5) > div > form > button'))
            )
            submit2.click()
        except:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(6) > div > form > button'))
            )
            submit2.click()
        
        # Fill form fields
        name = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[1]/div/input'))
        )
        name.send_keys(author_name)
        
        email = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[2]/div/input'))
        )
        email.clear()
        email.send_keys(author_email)
        
       
        
        number = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[4]/div/input'))
        )
        number.clear()
        number.send_keys(phone_number)
        
        ComName = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="archivnmfield"]'))
        )
        ComName.clear()
        ComName.send_keys("Coherent Market Insights")
        
        s1 = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="popup-archiv"]/div/a[1]'))
        )
        s1.click()
        
        # Handle category selection with better error handling
        Category_element = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
        )
        
        # Wait a moment for the dropdown to be fully loaded
        time.sleep(1)
        
        # Get all available options from the dropdown
        select_obj = Select(Category_element)
        available_options = [option.text.strip() for option in select_obj.options]
        log_to_status(f"Available dropdown options: {available_options}")
        
        # Use the category directly from the row data
        website_category = category.strip()
        log_to_status(f"Trying to select category: '{website_category}'")
        
        # Check if the exact category exists in the dropdown
        if website_category in available_options:
            log_to_status(f"Category '{website_category}' found in dropdown options")
        else:
            log_to_status(f"Category '{website_category}' NOT found in dropdown options")
            log_to_status(f"Available options are: {available_options}")
        
        try:
            # Try different selection methods
            select_obj.select_by_visible_text(website_category)
            log_to_status(f"Successfully selected category: '{website_category}'")
        except Exception as e:
            log_to_status(f"Error selecting category '{website_category}' by visible text: {e}")
            
            # Try selecting by value if visible text fails
            try:
                for option in select_obj.options:
                    if option.text.strip() == website_category:
                        select_obj.select_by_value(option.get_attribute('value'))
                        log_to_status(f"Successfully selected category by value: '{website_category}'")
                        break
                else:
                    raise Exception(f"Could not find option with text '{website_category}'")
            except Exception as e2:
                log_to_status(f"Error selecting category by value: {e2}")
                
                # Final fallback - try to select "Health & Medicine" directly if we have a health-related category
                try:
                    if "health" in website_category.lower() or "medicine" in website_category.lower():
                        select_obj.select_by_visible_text("Business,Economy,Finance,Banking & Insurance")
                        log_to_status("Selected 'Business,Economy,Finance,Banking & Insurance' as fallback for health-related category")
                    else:
                        select_obj.select_by_index(1)  # Select first real option
                        log_to_status("Selected first available option as final fallback")
                except Exception as e3:
                    log_to_status(f"Final fallback also failed: {e3}")
                    select_obj.select_by_index(0)  # Select any option to continue
        
        title = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[7]/div/input'))
        )
        
        title.clear()
        title.send_keys(article_title,)
        
        text = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="inhalt"]'))
        )
        text.clear()
        text.send_keys(multiline_text)
        
        about = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[9]/div/textarea'))
        )
        about.clear()
        multi = """Contact Us:

            Mr. Shah
            Coherent Market Insights
            533 Airport Boulevard,
            Suite 400, Burlingame,
            CA 94010, United States
            US: + 12524771362
            UK: +442039578553
            AUS: +61-8-7924-7805
            India: +91-848-285-0837"""
        about.send_keys(multi)
        address = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[10]/div/textarea'))
        )
        address.clear()
        random_author = random.choice(AUTHOR_DESCRIPTIONS)
        address_content = f"{random_author}\n\nAbout Us:\nCoherent Market Insights leads into data and analytics, audience measurement, consumer behaviors, and market trend analysis. From shorter dispatch to in-depth insights, CMI has exceled in offering research, analytics, and consumer-focused shifts for nearly a decade. With cutting-edge syndicated tools and custom-made research services, we empower businesses to move in the direction of growth. We are multifunctional in our work scope and have 450+ seasoned consultants, analysts, and researchers across 26+ industries spread out in 32+ countries."

        address.send_keys(address_content)
        
        image = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="bild"]'))
        )
        image.clear()
        image.send_keys(image_path)
        
        caption = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[12]/div/input'))
        )
        caption.clear()
        caption.send_keys("This is a test caption for the image.")
        
        notes = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[13]/div/textarea'))
        )
        notes.clear()
        notes.send_keys("This is a test notes section for the press release submission.")
        
        # Agree to terms
        tick1 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-agb"]'))
        )
        tick1.click()
        
        tick2 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-ds"]'))
        )
        tick2.click()
        
        # Submit form
        final = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/button'))
        )
        final.click()
        
        time.sleep(5)
        log_to_status(f"Selenium automation completed successfully for: {market_name}")
        driver.quit()  # Quit the browser tab after publication
        return True, market_name
        
    except Exception as e:
        market_name = row_data.get('Market Name', 'Unknown') if 'row_data' in locals() else 'Unknown'
        log_to_status(f"Selenium automation error for {market_name}: {e}")
        try:
            driver.quit()
        except:
            pass
        return False, market_name


def run_selenium_automation_all_rows(article_code, author_name, author_email, company_name, phone_number):
    """Run Selenium automation for all rows in the Excel file"""
    try:
        # Read all data from Excel
        import pandas as pd
        excel_path = r'C:\Users\akshat\Desktop\RPA\ROB.xlsx'
        df = pd.read_excel(excel_path)
        
        log_to_status(f"Found {len(df)} rows in Excel file")
        
        # Results tracking
        successful_submissions = []
        failed_submissions = []
        
        # Process each row
        for index, row in df.iterrows():
            log_to_status(f"\n{'='*50}")
            log_to_status(f"Processing Row {index + 1} of {len(df)}")
            log_to_status(f"{'='*50}")
            
            try:
                category = row['Category'] if 'Category' in row else ''
                success, market_name = run_selenium_automation_single(
                    row_data=row,
                    category=category,
                    article_code=article_code,
                    author_name=author_name,
                    author_email=author_email,
                    company_name=company_name,
                    phone_number=phone_number,
                    image_path=image_path if 'image_path' in locals() else None
                )
                
                if success:
                    successful_submissions.append(market_name)
                    log_to_status(f"✅ Successfully submitted: {market_name}")
                else:
                    failed_submissions.append(market_name)
                    log_to_status(f"❌ Failed to submit: {market_name}")
                
                # Add a delay between submissions to avoid overwhelming the server
                if index < len(df) - 1:  # Don't wait after the last submission
                    log_to_status("Waiting 360 seconds before next submission...")
                    time.sleep(20)
                    
            except Exception as e:
                market_name = row.get('Market Name', f'Row {index + 1}')
                failed_submissions.append(market_name)
                log_to_status(f"❌ Error processing row {index + 1} ({market_name}): {e}")
                continue
        
        # Final summary
        log_to_status(f"\n{'='*50}")
        log_to_status("FINAL SUMMARY")
        log_to_status(f"{'='*50}")
        log_to_status(f"Total rows processed: {len(df)}")
        log_to_status(f"Successful submissions: {len(successful_submissions)}")
        log_to_status(f"Failed submissions: {len(failed_submissions)}")
        
        if successful_submissions:
            log_to_status(f"\n✅ Successfully submitted:")
            for market in successful_submissions:
                log_to_status(f"  - {market}")
        
        if failed_submissions:
            log_to_status(f"\n❌ Failed submissions:")
            for market in failed_submissions:
                log_to_status(f"  - {market}")
        
        return len(successful_submissions), len(failed_submissions)
        
    except Exception as e:
        log_to_status(f"Error in run_selenium_automation_all_rows: {e}")
        return 0, 0

def run_selenium_automation(article_code, article_title, multiline_text, category, author_name, 
                          author_email, company_name, phone_number, image_path):
    """Enhanced run_selenium_automation function that reads category from Excel"""
    try:
        import random
        AUTHOR_DESCRIPTIONS = [
            """Author of this marketing PR:\nRavina Pandya, Content Writer, has a strong foothold in the market research industry. She specializes in writing well-researched articles from different industries, including food and beverages, information and technology, healthcare, chemical and materials, etc.\n """,
            """ Author of this marketing PR :\nMoney Singh is a seasoned content writer with over four years of experience in the market research sector. Her expertise spans various industries, including food and beverages, biotechnology, chemical and materials, defense and aerospace, consumer goods, etc. \n""",
            """ Author of this marketing PR:\n\nAlice Mutum is a seasoned senior content editor at Coherent Market Insights, leveraging extensive expertise gained from Openpr her previous role as a content writer. With seven years in content development, Alice masterfully employs SEO best practices and cutting-edge digital marketing strategies to craft high-ranking, impactful content. As an editor, she meticulously ensures flawless grammar and punctuation, precise data accuracy, and perfect alignment with audience needs in every research report. Alice's dedication to excellence and her strategic approach to content make her an invaluable asset in the world of market insights.\n"""
        ]
        log_to_status("Starting Selenium automation...")
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        options.add_argument("--start-maximized")
        
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.openpr.com/')
        
        # Handle cookie consent
        try:
            reject = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="cmpbntnotxt"]'))
            )
            reject.click()
        except:
            pass
        
        # Navigate to submit page
        submit = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="navbarText"]/ul/li[3]/a'))
        )
        submit.click()
        
        # Enter article code
        input_box = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="code"]'))
        )
        input_box.clear()
        input_box.send_keys(article_code)
        
        # Submit code
        try:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(5) > div > form > button'))
            )
            submit2.click()
        except:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(6) > div > form > button'))
            )
            submit2.click()
        
        # Fill form fields
        name = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[1]/div/input'))
        )
        name.send_keys(author_name)
        
        email = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[2]/div/input'))
        )
        email.clear()
        email.send_keys(author_email)
        
        '''pr_agency = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[3]/div/input'))
        )
        pr_agency.clear()
        pr_agency.send_keys(author_name)'''
        
        number = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[4]/div/input'))
        )
        number.clear()
        number.send_keys(phone_number)
        
        ComName = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="archivnmfield"]'))
        )
        ComName.clear()
        ComName.send_keys("Coherent Market Insights")
        
        s1 = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="popup-archiv"]/div/a[1]'))
        )
        s1.click()
        
        # Handle category selection with provided category argument
        Category_element = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
        )
        time.sleep(1)
        select_obj = Select(Category_element)
        available_options = [option.text.strip() for option in select_obj.options]
        log_to_status(f"Available dropdown options: {available_options}")
        website_category = category.strip()
        log_to_status(f"Trying to select category: '{website_category}'")
        try:
            select_obj.select_by_visible_text(website_category)
            log_to_status(f"Successfully selected category: '{website_category}'")
        except Exception as e:
            log_to_status(f"Error selecting category '{website_category}': {e}")
            try:
                select_obj.select_by_visible_text("Business,Economy,Finance,Banking & Insurance")
                log_to_status("Selected 'Business,Economy,Finance,Banking & Insurance'")
            except:
                select_obj.select_by_index(1)
                log_to_status("Selected first available option as final fallback")
        
        title = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[7]/div/input'))
        )
        title.clear()
        log_to_status(f"Trying to '{article_title}'")
        title.send_keys(article_title)
        
        text = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="inhalt"]'))
        )
        text.clear()
        text.send_keys(multiline_text)
        
        about = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[9]/div/textarea'))
        )
        about.clear()
        multi = """Contact Us:

Mr. Shah
Coherent Market Insights
533 Airport Boulevard,
Suite 400, Burlingame,
CA 94010, United States
US: + 12524771362
UK: +442039578553
AUS: +61-8-7924-7805
India: +91-848-285-0837
"""
        about.send_keys(multi)
        
        address = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[10]/div/textarea'))
        )
        address.clear()
        random_author = random.choice(AUTHOR_DESCRIPTIONS)
        address_content = f"{random_author}\n\nAbout Us:\nCoherent Market Insights leads into data and analytics, audience measurement, consumer behaviors, and market trend analysis. From shorter dispatch to in-depth insights, CMI has exceled in offering research, analytics, and consumer-focused shifts for nearly a decade. With cutting-edge syndicated tools and custom-made research services, we empower businesses to move in the direction of growth. We are multifunctional in our work scope and have 450+ seasoned consultants, analysts, and researchers across 26+ industries spread out in 32+ countries."

        address.send_keys(address_content)
        
        image = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="bild"]'))
        )
        image.clear()
        image.send_keys(image_path)
        
        caption = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[12]/div/input'))
        )
        caption.clear()
        caption.send_keys("This is a test caption for the image.")
        
        notes = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[13]/div/textarea'))
        )
        notes.clear()
        notes.send_keys("This is a test notes section for the press release submission.")
        
        # Agree to terms
        tick1 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-agb"]'))
        )
        tick1.click()
        
        tick2 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-ds"]'))
        )
        tick2.click()
        
        # Submit form
        final = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/button'))
        )
        final.click()
        
       
        
        time.sleep(10)
        log_to_status("Selenium automation completed successfully")
        return True
        
        
    except Exception as e:
        log_to_status(f"Selenium automation error: {e}")
        try:
            driver.quit()
        except:
            pass
        return False

def process_documents_auto_with_feedback(folder_path, article_code, author_name, author_email, company_name, phone_number,image_path):
    """Process documents automatically with status feedback"""
    global processing_status
    
    try:
        log_to_status(f"Starting auto processing. Folder: {folder_path}")
        
        excel_path = r'C:\Users\akshat\Desktop\RPA\ROB.xlsx'
        import random
        
        # Load Excel file
        log_to_status("Loading Excel file...")
        keywords_df = pd.read_excel(excel_path)
        market_names = keywords_df['Market Name'].dropna().tolist()
        
        processing_status['total'] = len(market_names)
        log_to_status(f"Found {len(market_names)} market names to process")
        
        processed_count = 0
        TITLE_PROMPTS = [
    "Is Booming Worldwide 2025-2032",
    "Generated Opportunities, Future Scope 2025-2032",
    "Future Business Opportunities 2025-2032",
    "Growth in Future Scope 2025-2032",
    "Is Booming So Rapidly Growth by 2032",
    "Is Booming So Rapidly with CAGR of 6.9%",
    "An Analysis of Size, Shares, Business Growth, and Upcoming Trends Forecast 2025-2032",
    "2025-2032 Emerging Trends in Industry Dynamics, Size Insights, Share, and Future Growth",
    "Set to Witness Significant Growth by 2025-2032",
    "to Witness Massive Growth by 2032",
    "Size, Share 2025 Analysis of Rising Business Opportunities with Prominent Investment, Forecast to 2032",
    "by Trends, Dynamic Innovation in Technology and 2032 Forecast, Opportunities, and Challenges, Trends",
    "Size 2025 Emerging Demands, Share, Trends, Futuristic Opportunity, Share and Forecast To 2032",
    "Size, Share Growth Status, Emerging Technology, Key Players, Industry Challenges, and Forecast till 2032",
    "Insights 2025-2032: Global Expansion, Revenue Trends, and Strategic Growth Plans",
    "Outlook 2025-2032: Growth Drivers, Share, And Trends As Revealed In New Report"
]

        
        for i, market_name in enumerate(market_names):
            processing_status['progress'] = i
            processing_status['current_file'] = market_name
            processing_status['message'] = f"Auto-processing {i+1} of {len(market_names)}: {market_name}"
            
            doc_file = f"ROB_{market_name}.doc"
            doc_path = os.path.join(folder_path, doc_file)
            
            log_to_status(f"Looking for file: {doc_path}")
            
            if os.path.exists(doc_path):
                log_to_status(f"Processing: {market_name}")
                
                # Extract text
                processing_status['message'] = f"Extracting text from {market_name}..."
                multiline_text = text_of_press_release(doc_path)
                # Get companies for this market from Excel row
                matching_row = keywords_df[keywords_df['Market Name'] == market_name]
                if not matching_row.empty:
                    companies = matching_row.iloc[0].get('Companies covered', '')
                else:
                    companies = ''
                if companies and isinstance(companies, str) and companies.strip():
                    company_list = [c.strip() for c in companies.split(',') if c.strip()]
                    first_five_companies = ', '.join(company_list[:3]) if company_list else 'No companies available'
                else:
                    first_five_companies = 'No companies available'
                random_prompt = random.choice(TITLE_PROMPTS)
                x = f"{market_name} {random_prompt}|{first_five_companies}"

                # Refine article_title using OpenAI for grammar correction
                def refine_title_with_openai(title):
                    try:
                        import openai
                        client = openai.OpenAI(api_key=OPENAI_CONFIG['API_KEY'])
                        prompt = f"Please improve the grammar, structure, and readability of this press release title to make it more interesting and engaging for readers. Keep all original words intact - only rearrange, or adjust formatting as needed and there should be no parenthesis at end or start of title : '{title}'" 
                        response = client.chat.completions.create(
                            model=OPENAI_CONFIG['MODEL'],
                            messages=[
                                {"role": "system", "content": "You are an expert editor for press releases."},
                                {"role": "user", "content": prompt}
                            ],
                            max_tokens=60,
                            temperature=2.0
                        )
                        return response.choices[0].message.content.strip()
                    except Exception as e:
                        log_to_status(f"OpenAI title refinement error: {e}")
                        return title

                article_title = refine_title_with_openai(x)
                
                category = matching_row.iloc[0].get('Category', '') if not matching_row.empty else ''
                # Run automation
                processing_status['message'] = f"Submitting {market_name} via automation..."
                success = run_selenium_automation(article_code,article_title, multiline_text, category, 
                                                author_name, author_email, company_name, phone_number, image_path)
                
                if success:
                    log_to_status(f"SUCCESS: Published {market_name}")
                    processed_count += 1
                else:
                    log_to_status(f"FAILED: Could not publish {market_name}")
                
                time.sleep(10)
                  # Wait 60 seconds between submissions to avoid rate limiting
                
            else:
                log_to_status(f"ERROR: File not found: {doc_path}")
        
        processing_status['active'] = False
        processing_status['message'] = f"Auto-processing complete! Published {processed_count} of {len(market_names)} articles"
        log_to_status(f"Auto processing complete. Published {processed_count} articles.")
        
    except Exception as e:
        processing_status['active'] = False
        processing_status['message'] = f"Error: {str(e)}"
        log_to_status(f"EXCEPTION: Auto processing error: {e}")

def process_documents_manual_with_feedback(folder_path, article_code, author_name, author_email, company_name, phone_number,image_path):
    """Process documents with manual intervention and status feedback"""
    global processing_status
    import random
    
    try:
        log_to_status(f"Starting manual processing. Folder: {folder_path}")
        
        excel_path = r'C:\Users\akshat\Desktop\RPA\ROB.xlsx'
        
        # Load Excel file
        log_to_status("Loading Excel file...")
        keywords_df = pd.read_excel(excel_path)
        market_names = keywords_df['Market Name'].dropna().tolist()
        
        processing_status['total'] = len(market_names)
        log_to_status(f"Found {len(market_names)} market names to process")
        
        processed_count = 0

        TITLE_PROMPTS = [
    "Is Booming Worldwide 2025-2032",
    "Generated Opportunities, Future Scope 2025-2032",
    "Future Business Opportunities 2025-2032",
    "Growth in Future Scope 2025-2032",
    "Is Booming So Rapidly Growth by 2032",
    "Is Booming So Rapidly with CAGR of 6.9%",
    "An Analysis of Size, Shares, Business Growth, and Upcoming Trends Forecast 2025-2032",
    "2025-2032 Emerging Trends in Industry Dynamics, Size Insights, Share, and Future Growth",
    "Set to Witness Significant Growth by 2025-2032",
    "to Witness Massive Growth by 2032",
    "Size, Share 2025 Analysis of Rising Business Opportunities with Prominent Investment, Forecast to 2032",
    "by Trends, Dynamic Innovation in Technology and 2032 Forecast, Opportunities, and Challenges, Trends",
    "Size 2025 Emerging Demands, Share, Trends, Futuristic Opportunity, Share and Forecast To 2032",
    "Size, Share Growth Status, Emerging Technology, Key Players, Industry Challenges, and Forecast till 2032",
    "Insights 2025-2032: Global Expansion, Revenue Trends, and Strategic Growth Plans",
    "Outlook 2025-2032: Growth Drivers, Share, And Trends As Revealed In New Report"
]

        
        for i, market_name in enumerate(market_names):
            processing_status['progress'] = i
            processing_status['current_file'] = market_name
            processing_status['message'] = f"Processing {i+1} of {len(market_names)}: {market_name}"
            
            doc_file = f"ROB_{market_name}.doc"
            doc_path = os.path.join(folder_path, doc_file)
            
            log_to_status(f"Looking for file: {doc_path}")
            
            if os.path.exists(doc_path):
                log_to_status(f"Processing: {market_name}")
                
                # Extract text
                processing_status['message'] = f"Extracting text from {market_name}..."
                multiline_text = text_of_press_release(doc_path)
                # Get companies for this market from Excel row
                matching_row = keywords_df[keywords_df['Market Name'] == market_name]
                if not matching_row.empty:
                    companies = matching_row.iloc[0].get('Companies covered', '')
                else:
                    companies = ''
                if companies and isinstance(companies, str) and companies.strip():
                    company_list = [c.strip() for c in companies.split(',') if c.strip()]
                    first_five_companies = ', '.join(company_list[:5]) if company_list else 'No companies available'
                else:
                    first_five_companies = 'No companies available'
                random_prompt = random.choice(TITLE_PROMPTS)
                article_title = f"{market_name} {random_prompt} {first_five_companies}"
                
                # Get category for this market from Excel row
                category = matching_row.iloc[0].get('Category', '') if not matching_row.empty else ''
                # Run automation
                processing_status['message'] = f"Submitting {market_name} via automation..."
                success = run_selenium_automation(article_code, article_title, multiline_text, category, 
                                                author_name, author_email, company_name, phone_number, image_path)
                
                if success:
                    log_to_status(f"Published {market_name}")
                    processed_count += 1
                
                time.sleep(5)
                
            else:
                log_to_status(f"ERROR: File not found: {doc_path}")
        
        processing_status['active'] = False
        processing_status['message'] = f"Processing complete! Published {processed_count} of {len(market_names)} articles"
        log_to_status(f"Manual processing complete. Published {processed_count} articles.")
        
    except Exception as e:
        processing_status['active'] = False
        processing_status['message'] = f"Error: {str(e)}"
        log_to_status(f"EXCEPTION: Manual processing error: {e}")

# ============================================================================
# ROB PROCESSING ROUTES
# ============================================================================

@app.route('/rob', methods=['GET', 'POST'])
def rob():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        open_pr_id = request.form.get('open_pr_id')
        mobile = request.form.get('mobile')
        extract_count = int(request.form.get('extract_count', 200))

        # Validate required fields
        if not all([username, email, open_pr_id, mobile, extract_count]):
            flash('All fields are required!')
            return redirect(request.url)

        file = request.files.get('file')
        if not file or file.filename == '':
            flash('Excel file is required!')
            return redirect(request.url)

        if not allowed_file(file.filename):
            flash('Only Excel files (.xlsx, .xls) and CSV files are allowed!')
            return redirect(request.url)

        # Use secure_filename to avoid path issues
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)

        # Store user data in session for later use
        session['username'] = username
        session['email'] = email
        session['open_pr_id'] = open_pr_id
        session['mobile'] = mobile
        
        return redirect(url_for('process_rob', file_path=input_path,
                                username=username, email=email,
                                open_pr_id=open_pr_id, mobile=mobile,
                                extract_count=extract_count))
    return render_template('rob.html')

@app.route('/process_rob')
def process_rob():
    file_path = request.args.get('file_path')
    username = request.args.get('username')
    email = request.args.get('email')
    open_pr_id = request.args.get('open_pr_id')
    mobile = request.args.get('mobile')
    extract_count = int(request.args.get('extract_count', 200))

    if not file_path or not os.path.exists(file_path):
        flash('Missing or invalid file path')
        return redirect(url_for('rob'))

    try:
        # Read the cleaned ROB file
        if file_path.endswith('.csv'):
            df_original = pd.read_csv(file_path)
        else:
            df_original = pd.read_excel(file_path, engine='openpyxl')

        total_rows = len(df_original)
        
        if total_rows < extract_count:
            flash(f'⚠️ File only has {total_rows} rows, but you requested {extract_count} rows!')
            extract_count = total_rows

        # Step 1: Extract top N rows
        extracted_rows = df_original.head(extract_count).copy()
        
        # Step 2: Get remaining rows (original minus extracted)
        remaining_rows = df_original.iloc[extract_count:].copy()

        # Step 3: Create timestamp for remaining file
        today = datetime.today()
        timestamp = f"{today.year}_{today.month:02d}_{today.day:02d}"
        
        # Step 4: Save extracted rows as ROB.xlsx to Desktop/RPA
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        rob_output_path = os.path.join(rpa_folder, "ROB.xlsx")
        extracted_rows.to_excel(rob_output_path, index=False)
        
      # Step 5: Save remaining rows with timestamp in date-wise folder under Weekly_RID
        weekly_rid_folder = os.path.join(r"C:\Users\akshat\Desktop\RPA\Weekly_RID", str(today.year), f"{today.month:02d}", f"{today.day:02d}")
        os.makedirs(weekly_rid_folder, exist_ok=True)
        remaining_filename = f"cleaned_weekly_rid_{timestamp}.xlsx"
        remaining_output_path = os.path.join(weekly_rid_folder, remaining_filename)
        remaining_rows.to_excel(remaining_output_path, index=False)
        
        # Step 6: Store info in session for the result page
        session['rob_file_path'] = rob_output_path
        session['remaining_file_path'] = remaining_output_path
        session['remaining_filename'] = remaining_filename
        session['extracted_count'] = extract_count
        session['remaining_count'] = len(remaining_rows)
        session['total_count'] = total_rows
        
        flash(f'✅ Successfully processed {total_rows} rows!')
        flash(f'📁 Remaining {len(remaining_rows)} rows saved in: {weekly_rid_folder} as {remaining_filename} (ready for download)')
        
        # Use render_template instead of redirect
        return render_template('rob_result.html',
                             extracted_count=extract_count,
                             remaining_count=len(remaining_rows),
                             total_count=total_rows,
                             username=username,
                             records_processed=total_rows,
                             weekly_rid_folder=weekly_rid_folder,
                             remaining_filename=remaining_filename)

    except Exception as e:
        flash(f'❌ Error processing ROB file: {str(e)}')
        return redirect(url_for('rob'))

@app.route('/download_remaining_rob')
def download_remaining_rob():
    """Download the remaining ROB file (original minus extracted rows)"""
    try:
        remaining_file_path = session.get('remaining_file_path')
        remaining_filename = session.get('remaining_filename', 'cleaned_rob_remaining.xlsx')
        
        if remaining_file_path and os.path.exists(remaining_file_path):
            return send_file(remaining_file_path, as_attachment=True, download_name=remaining_filename)
        else:
            flash('❌ Remaining ROB file not found. Please process a file first.')
            return redirect(url_for('rob'))
    except Exception as e:
        flash(f'❌ Error downloading remaining file: {str(e)}')
        return redirect(url_for('rob'))

@app.route('/download_extracted_rob')
def download_extracted_rob():
    """Download the extracted ROB.xlsx file and trigger OpenAI content generation"""
    try:
        rob_file_path = session.get('rob_file_path')
        
        if rob_file_path and os.path.exists(rob_file_path):
            
            # Start OpenAI content generation in background thread with 5-second delay
            print("🔄 Starting background OpenAI content generation...")
            threading.Thread(target=delayed_openai_content_generation, args=(5,)).start()
            
            return send_file(rob_file_path, as_attachment=True, download_name='ROB.xlsx')
        else:
            flash('❌ ROB.xlsx file not found. Please process a file first.')
            return redirect(url_for('rob'))
            
    except Exception as e:
        flash(f'❌ Error downloading ROB file: {str(e)}')
        return redirect(url_for('rob'))


def delayed_openai_content_generation(delay_seconds=5):
    """Generate content using OpenAI after delay (replaces Power Automate)"""
    try:
        print(f"⏳ Waiting {delay_seconds} seconds before starting content generation...")
        time.sleep(delay_seconds)
        
        print("🤖 Starting OpenAI content generation...")
        
        # ROB file path
        rob_file_path = r"C:\Users\akshat\Desktop\RPA\ROB.xlsx"
        
        if not os.path.exists(rob_file_path):
            print("❌ ROB.xlsx file not found!")
            return
        
        # Check if API key is configured
        if OPENAI_CONFIG['API_KEY'] == 'your-openai-api-key-here':
            print("❌ OpenAI API key not configured! Using fallback content.")
        
        # Read ROB file
        df = pd.read_excel(rob_file_path)
        print(f"📊 Found {len(df)} markets in ROB file")
        
        # Create output directory with current date
        today = datetime.today()
        output_dir = os.path.join(
            r"C:\Users\akshat\Desktop\RPA\Files",
            str(today.year),
            f"{today.month:02d}",
            f"{today.day:02d}"
        )
        os.makedirs(output_dir, exist_ok=True)
        print(f"📁 Output directory: {output_dir}")
        
        successful = 0
        failed = 0
        
        # Process each row
        for index, row in df.iterrows():
            try:
                # Extract market data
                
                #print(f"[{index+1}/{len(df)}] Processing: {data['market_name']}")
                
                # Generate content
                content = generate_blog_from_row(row)
                
                # Save document
                success, filepath = save_market_document(row.get('Market Name'), content, output_dir)
                
                if success:
                    print(f"✅ Generated: {os.path.basename(filepath)}")
                    successful += 1
                else:
                    print(f"❌ Failed to save: (row.get('Market Name')")
                    failed += 1
                
                # Rate limiting
                time.sleep(2)
                
            except Exception as e:
                print(f"❌ Error processing  {e}")
                failed += 1
                continue
        
        print(f"✅ Content generation completed!")
        print(f"📊 Successful: {successful}, Failed: {failed}")
        print(f"📁 Files saved in: {output_dir}")
        
    except Exception as e:
        print(f"❌ Error in content generation: {e}")


# ============================================================================
# ADD THESE SIMPLE HELPER FUNCTIONS
# ============================================================================

import re
import openai

# OpenAI config
OPENAI_CONFIG = {
    'API_KEY': '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',
    'MODEL': 'gpt-4.1-mini',  # Or 'gpt-4o' or 'gpt-3.5-turbo'
    'TEMPERATURE': 0.7
}

def generate_blog_from_row(row):
    """Extract data from row, format prompt, and generate blog via OpenAI."""
    try:
        # Extract values
        market_name = row.get('Market Name', 'Unknown Market')
        forecast_period = row.get('Forecast Period', 'Not specified')
        market_size_year = row.get('Market Size Year', 'Not specified')
        market_size = row.get('Market Size', '')
        cagr = row.get('CAGR', 'Not specified')
        key_players = row.get('Key Players', 'Not specified')

        # Extract size values from combined string
        match = re.search(r'USD ([\d.]+ Bn); Market Size in 2032: USD ([\d.]+ Bn)', market_size)
        if match:
            market_size_2025 = f"USD {match.group(1)}"
            market_size_2032 = f"USD {match.group(2)}"
        else:
            market_size_2025 = ""
            market_size_2032 = ""

        # Prepare prompt string from extracted values
        output = ( f"""
We are coherent market insights and going to publish this data on third part PR website, thus do not mention any competitors name (market intelligence firm, market research, consulting and data and analytics companies), and take ownership for market share, size, growth rate as this research done by our own research team however any real instances and stats you can use reference such as company or organisation or others as mentioned in detail explanation. Write the blog covering the below pointers: Start content with market size and overview paragraph (do not add heading and do not use bullet for this section), ➔ add first CTA link here, ➤Actionable Insights, ➤Market Segment and Regional coverage, ➔ add Second CTA link here, ➤Key players, ➤Growth factors, ➔add Third CTA link here, ➤ Market trends, ➤Market Opportunities, and ❓ Frequently Asked Questions, All these pointers should act as a heading to respective paragraphs, do not miss any bullet foe above given. Consider the following points while generating content: Flow of information, all given secondary keywords must be covered, and there should be a heading for each paragraph or bullet pointers. I need this content to be very presentable format thus make sure there should be space after and before of each heading or CTA links (first CTA link, Actionable Insights, Market Segment and Regional coverage, Second CTA link, Key players, Growth factors, Third CTA link, Market trends, Market Opportunities, and Frequently Asked Questions), this will increase the readibility. Cover content in in bullet pointers whenever possible each paragraph should be short, about 100 to 120 words. Our readers are already experts in the field, so always try to generate content that provides unique insights and value addition for experts. Thus, while generating each piece of content, it should be data-backed with actual instances and stats from recent years 2025 and 2024, covering maximum stats that increase the authenticity and are enough to support or make decisions based upon reading this blog. Do not add generic content that is already known to readers, such as definitions, advantages, disadvantages, or other generic content. Flow of information should be as below. Start the blog with writing 5 to 7 sentence ling paragraph start content with talking about Industry in 2 to 3 sentences and should be matches with heading of the blog. followed by 2 to 4 sentence on market size and CAGR formatted as The Global (Insert Market Name) Market size is estimated to be valued at USD (Insert Market Value for 2025) in 2025 (market size) and is expected to reach USD (Insert Market Value for 2032) by 2032 (market forecast), exhibiting a compound annual growth rate (CAGR) of (Insert CAGR)% from 2025 to 2032. Do not add heading and do not use bullet for this section. (Make sure all values such as market size, CAGR, key players should reflect exactly same in output as provided in input). Then First CTA link. Then Actionable Insights: In Actionable Insights, cover 3 to 4 actionable insights in 10 to 12 sentences each insights comprises of 2 t 3 sentences covering supporting data for each insights. Insights could be essential quantitative indicators that lead to market size. For example, supply-side indicators (ex. Production Capacity, Pricing, and Exports. but not necessary to add if not matches to given market), demand-side indicators (ex. Pricing, Imports, Various Use Cases across industries. but not necessary to add if not matches to given market), micro-indicators, nano-size indicators—focus more on the quantitative aspects. Each actionable insights must have two sentence stats or actual instance examples from the recent year to support each point given in actionable insights, so that each given point look complete and meaningful. Next part is Market segment and regional Coverage where enlist the all subsegment under each segment categories and fragment region into given format. Comprehensive Segmentation and Classification of the Report: » By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. » By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. » By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. Regional and Country Analysis: » North America: U.S. and Canada » Latin America: Brazil, Argentina, Mexico, and Rest of Latin America » Europe: Germany, U.K., Spain, France, Italy, Benelux, Denmark, Norway, Sweden, Russia, and Rest of Europe » Asia Pacific: China, Taiwan, India, Japan, South Korea, Indonesia, Malaysia, Philippines, Singapore, Australia, and Rest of Asia Pacific » Middle East & Africa: Bahrain, Kuwait, Oman, Qatar, Saudi Arabia, United Arab Emirates, Israel, South Africa, North Africa, Central Africa, and Rest of MEA. Then Second CTA link. Then key Players: List 12 to 20 highly relevant key players for the given market. Furthermore, add 2-3 statements on competitive strategies adopted by a few key players, mentioning actual strategies and entities involved along with the actual outcome. Growth Factors: Growth factor heading and short description with supporting stats or examples from the recent year in the content. Then Add Third CTA link. Then Market Trends: Market Trend heading and short description with supporting stats or examples from the recent year in the content. Then Market Opportunities: Provide 3 to 4 market opportunities, 2-3 opportunities based upon segment and one opportunity based upon region. Each opportunity described in two to three sentences and supported by actual data. For each opportunity, identify a single segment and provide a short description of the opportunity within that segment. Similarly, highlight one region, along with a brief description of the opportunity in each regional market. Make sure to not to mention our organization name or relevant terms anywhere in the output such as coherent market insights or our analyst team or our research team Given
 Market Name and Data:f"Market Name- {market_name}; "
            f"CAGR:{cagr}; "
            f"Forecast period is: {forecast_period}; "
            f"Market Size for {market_size_year} is {market_size_2025}; "
            f"Market Size in 2032: {market_size_2032}; "
            f"Key players: {key_players}"
Key Reasons for Buying the (insert market name here) Report: ✦ Comprehensive analysis of the changing competitive landscape ✦ Assists in decision-making processes for the businesses along with detailed strategic planning methodologies ✦ The report offers forecast data and an assessment of the (insert market name here) ✦ Helps in understanding the key product segments and their estimated growth rate ✦ In-depth analysis of market drivers, restraints, trends, and opportunities ✦ Comprehensive regional analysis of the (insert market name here) ✦ Extensive profiling of the key stakeholders of the business sphere ✦ Detailed analysis of the factors influencing the growth of the (insert market name here). From an SEO perspective, we need to cover all given keywords from the list below. However, they should appear naturally so that the content flow looks natural for the reader. Keyword List: market share, market size, market research, market insights, market trends, market opportunities, market challenges, market growth, market forecast, market companies, market players, market analysis, market drivers, market restraints, market scope, market dynamics, market segments, market report, market growth strategies, market revenue, industry size, industry share, industry trends, and business growth, furthermore - Market size and market report, market revenue, market share, trends keywords are mandatory to be added twice in content. In addition to the above requirement, in 5 places, add the actual market name along with the above keywords so that long-tail keywords will be generated. These long-tail keywords are market name + size, market name + report, market name + revenue, market name + share, market name + trends. Make sure all given keywords are naturally fit, do not try to infuse forcefully, flow of information should be natural and meaningful, furthermore make sure spelling and structure of sentences from generated output are grammatically correct. Furthermore, based on the market name, create a set of Frequently Asked Questions that are highly relevant and customized to the specific market. The sample Frequently Asked Questions below are for understanding purposes only. For the given market, questions can be completely replaced. However, please tailor the actual questions to the market name and the insights provided in the report: 1. Who are the dominant players in the (Market Name) market? 2. What will be the size of the (Market Name) market in the coming years? 3. Which end users industry has the largest growth opportunity? 4. How will market development trends evolve over the next five years? 5. What is the nature of the competitive landscape and challenges in the (Market Name) market? 6. What go-to-market strategies are commonly adopted in the (Market Name) market? Make sure to answer to all FAQs. In the case of country-level markets, please exclude the word 'Global' and Market Opportunities where other regions are mentioned. Make sure to add catchy bullet in generated output. I have shared the reference bullet with you. Make sure to add this bullet. For heading use these bullet- ➤Actionable Insights, ➤Market Segment and Regional Coverage, ➔ Inserted Second CTA link, ➤Key Players, ➤Growth factors, ➔ Inserted Third CTA link, ➤ Market Trends, ➤Market Opportunities, and ❓ Frequently Asked Questions. Make sure do not miss any bullet including CTA bullet which is ➔. For subpointers under main headings use bullets which is in reference as provided- Actionable Insights ●, Market Segment and Regional Coverage● , Key players●, Growth Factors●,  Market Trends●, Market Opportunities●. Make sure to use these bullets for given subpointers. Ensure proper bullet formatting so that each point is marked with a single bullet only, and avoid placing two bullets adjacent to each other.
""")

        # Send to OpenAI
        client = openai.OpenAI(api_key=OPENAI_CONFIG['API_KEY'])

        response = client.chat.completions.create(
            model=OPENAI_CONFIG['MODEL'],
            messages=[
                {"role": "user", "content": output}
            ],
            temperature=OPENAI_CONFIG.get('TEMPERATURE', 0.7)
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"OpenAI error: {e}")
        return "Error generating content."

def save_market_document(market_name, content, output_folder):
    """Save content as Word document"""
    try:
        doc = Document()
        doc.add_heading(f"{market_name} - Market Research Report", level=1)
        
        # Add content paragraphs
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Add contact info
      
        
        # Save file
        safe_name = "".join(c for c in market_name if c.isalnum() or c in (' ', '_')).strip()
        filename = f"ROB_{safe_name}.doc"
        filepath = os.path.join(output_folder, filename)
        doc.save(filepath)
        
        return True, filepath
        
    except Exception as e:
        print(f"Error saving document: {e}")
        return False, None


'''app.route('/api/auto_trigger_power_automate', methods=['POST'])
def auto_trigger_power_automate():
    """API endpoint for auto-triggering Power Automate"""
    try:
        # Check if we should trigger (based on recent download)
        if session.get('trigger_power_automate'):
            # Clear the flag
            session['trigger_power_automate'] = False
            
            # Trigger in background
            threading.Thread(target=delayed_power_automate_trigger, args=(0,)).start()
            
            return jsonify({
                'status': 'success', 
                'message': 'Power Automate triggered automatically after ROB download'
            })
        else:
            return jsonify({
                'status': 'error', 
                'message': 'No recent ROB download detected'
            })
    except Exception as e:
        return jsonify({
            'status': 'error', 
            'message': f'Error: {str(e)}'
        })'''

# ============================================================================
# WEEKLY REPORT ROUTES
# ============================================================================

@app.route('/weekly_report', methods=['GET', 'POST'])
def weekly_report():
    if request.method == 'POST':
        form_type = request.form.get('form_type')
        
        if form_type == 'backend_processing':
            return handle_backend_processing()
        else:
            return handle_rid_analysis()
    
    # GET request - show form (no data to display)
    return render_template('weekly_report.html', qualified_rids=None, filter_summary=None, backend_result=None)

def handle_rid_analysis():
    """Handle RID analysis with dual file input - ranking sheet + cleaned ROB file"""
    try:
        print("RID Analysis POST request received!")
        
        # Get filter parameters from form
        min_search_volume = int(request.form.get('min_search_volume', 5000))
        competition_level = request.form.get('competition_level', 'Low')
        analyze_trends = request.form.get('analyze_trends') == 'on'
        
        print(f"User Filters: Search >= {min_search_volume}, Competition = {competition_level}")
        print(f"Google Trends: {'Enabled' if analyze_trends else 'Disabled'}")
        
        # Validate form inputs
        if not min_search_volume or min_search_volume < 0:
            flash('❌ Please enter a valid minimum search volume!')
            return redirect(request.url)
            
        if not competition_level:
            flash('❌ Please select a competition level!')
            return redirect(request.url)
        
        # Handle RANKING SHEET upload
        ranking_file = request.files.get('ranking_file')
        if not ranking_file or ranking_file.filename == '':
            flash('❌ Please select a ranking Excel file!')
            return redirect(request.url)

        if not allowed_file(ranking_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed for ranking sheet!')
            return redirect(request.url)

        # Handle CLEANED ROB FILE upload
        rob_file = request.files.get('cleaned_rob_file')
        if not rob_file or rob_file.filename == '':
            flash('❌ Please select a cleaned ROB Excel file!')
            return redirect(request.url)

        if not allowed_file(rob_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed for ROB file!')
            return redirect(request.url)

        # Save both uploaded files
        ranking_filename = secure_filename(ranking_file.filename)
        ranking_path = os.path.join(app.config['UPLOAD_FOLDER'], ranking_filename)
        ranking_file.save(ranking_path)
        print(f"Ranking file saved: {ranking_path}")
        
        rob_filename = secure_filename(rob_file.filename)
        rob_path = os.path.join(app.config['UPLOAD_FOLDER'], rob_filename)
        rob_file.save(rob_path)
        print(f"ROB file saved: {rob_path}")
        
        # Process both files and get qualified ROB data
        result_summary = process_dual_files_and_extract_rob(
            ranking_path, rob_path, min_search_volume, competition_level, analyze_trends
        )
        
        # Format success/warning messages based on results
        if result_summary['success']:
            flash(f'✅ Success! Found {result_summary["qualified_rids_count"]} qualified RIDs')
            flash(f'✅ Extracted {result_summary["matched_rob_rows"]} matching ROB rows')
            flash(f'📁 Weekly ROB.xlsx saved to Desktop/RPA folder!')
            print(f"Dual file processing completed: {result_summary}")
        else:
            flash(f'❌ Error: {result_summary.get("error", "Unknown error")}')
            result_summary = None
        
        # Clean up uploaded files after processing
        try:
            os.remove(ranking_path)
            os.remove(rob_path)
            print(f"Cleaned up uploaded files")
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up files: {cleanup_error}")
        
        # Render template with results
        return render_template('weekly_report.html', 
                              qualified_rids=result_summary.get('qualified_rids', []) if result_summary else [],
                              filter_summary=result_summary.get('filter_summary', {}) if result_summary else {},
                              backend_result=None,
                              rob_extraction_result=result_summary)
        
    except ValueError as ve:
        print(f"Value Error: {ve}")
        flash('❌ Invalid input values. Please check your filters.')
        return redirect(request.url)
    except Exception as e:
        print(f"Error: {e}")
        flash(f'❌ Error processing files: {str(e)}')
        return redirect(request.url)

def process_dual_files_and_extract_rob(ranking_path, rob_path, min_search_volume, competition_level, analyze_trends):
    """Process ranking sheet and ROB file together, extract matching rows"""
    try:
        print(f"\n=== PROCESSING DUAL FILES ===")
        print(f"Ranking file: {ranking_path}")
        print(f"ROB file: {rob_path}")
        
        # STEP 1: Process ranking sheet to get qualified RIDs
        print("\n📊 STEP 1: Processing ranking sheet...")
        qualified_rids, filter_summary, updated_ranking_path = get_qualified_rids_and_remove_trending(
            ranking_path, min_search_volume, competition_level, analyze_trends
        )
        
        if not qualified_rids:
            return {
                'success': False,
                'error': 'No qualified RIDs found in ranking sheet with your filter criteria'
            }
        
        print(f"✅ Found {len(qualified_rids)} qualified RIDs: {qualified_rids}")
        
        # STEP 2: Process ROB file and extract matching rows
        print(f"\n📋 STEP 2: Processing ROB file and extracting matching rows...")
        
        # Read the cleaned ROB file
        if rob_path.endswith('.csv'):
            rob_df = pd.read_csv(rob_path)
        else:
            rob_df = pd.read_excel(rob_path, engine='openpyxl')
        
        total_rob_rows = len(rob_df)
        print(f"ROB file loaded: {total_rob_rows} rows")
        print(f"ROB file columns: {list(rob_df.columns)}")
        
        # Find Report ID column
        report_id_column = None
        possible_columns = ['Report ID', 'ReportID', 'report_id', 'ID', 'Report_ID', 'Market Name']
        
        for col in possible_columns:
            if col in rob_df.columns:
                report_id_column = col
                break
        
        if not report_id_column:
            return {
                'success': False,
                'error': f'Report ID column not found in ROB file. Available columns: {list(rob_df.columns)}'
            }
        
        print(f"Using Report ID column: {report_id_column}")
        
        # Convert qualified_rids to same type as Report ID column
        rob_df[report_id_column] = rob_df[report_id_column].astype(str).str.strip()
        qualified_rids_str = [str(rid).strip() for rid in qualified_rids]
        
        print(f"Sample Report IDs in ROB file: {rob_df[report_id_column].head().tolist()}")
        print(f"Looking for RIDs: {qualified_rids_str}")
        
        # Filter ROB rows that match qualified RIDs
        matching_rob_rows = rob_df[rob_df[report_id_column].isin(qualified_rids_str)].copy()
        matched_count = len(matching_rob_rows)
        
        print(f"Found {matched_count} matching ROB rows")
        
        if matched_count == 0:
            return {
                'success': False,
                'error': f'No matching Report IDs found in ROB file. Check if Report IDs {qualified_rids} exist in the ROB file.'
            }
        
        # Show which RIDs were found and missing
        found_rids = matching_rob_rows[report_id_column].tolist()
        missing_rids = [rid for rid in qualified_rids_str if rid not in found_rids]
        
        print(f"Found Report IDs: {found_rids}")
        if missing_rids:
            print(f"Missing Report IDs: {missing_rids}")
        
        # STEP 3: Save to Desktop/RPA folder as ROB.xlsx
        print(f"\n💾 STEP 3: Saving to Desktop...")
        
        # Create RPA folder on Desktop if it doesn't exist
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
            print(f"Created RPA folder: {rpa_folder}")
        
        # Save as ROB.xlsx in Desktop/RPA folder
        output_path = os.path.join(rpa_folder, "weekly_RID.xlsx")
        
        # Use xlsxwriter for better performance
        with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
            matching_rob_rows.to_excel(writer, index=False, sheet_name='ROB_Data')
        
        print(f"✅ weekly_RID.xlsx saved to: {output_path}")
        
        # Display sample of extracted data
        print("\nSample of extracted ROB data:")
        print(matching_rob_rows.head(2))
        
        # Create comprehensive summary
        summary = {
            'success': True,
            'qualified_rids': qualified_rids,
            'qualified_rids_count': len(qualified_rids),
            'total_rob_rows': total_rob_rows,
            'matched_rob_rows': matched_count,
            'found_rids': found_rids,
            'missing_rids': missing_rids,
            'output_path': output_path,
            'report_id_column': report_id_column,
            'filter_summary': filter_summary
        }
        
        return summary
        
    except Exception as e:
        print(f"Error in dual file processing: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def handle_backend_processing():
    """Handle backend file processing form submission"""
    try:
        print("Backend Processing POST request received!")
        
        # Get processing options
        auto_detect_header = request.form.get('auto_detect_header') == 'on'
        clean_columns = request.form.get('clean_columns') == 'on'
        remove_empty_rows = request.form.get('remove_empty_rows') == 'on'
        
        print(f"Processing options: Header={auto_detect_header}, Clean={clean_columns}, Remove Empty={remove_empty_rows}")
        
        # Handle backend file upload
        backend_file = request.files.get('backend_file')
        if not backend_file or backend_file.filename == '':
            flash('❌ Please select a backend Excel file!')
            return redirect(request.url)

        if not allowed_backend_file(backend_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) are allowed for backend processing!')
            return redirect(request.url)

        # Save uploaded backend file
        backend_filename = secure_filename(backend_file.filename)
        backend_path = os.path.join(app.config['UPLOAD_FOLDER'], backend_filename)
        backend_file.save(backend_path)
        print(f"Backend file saved: {backend_path}")
        
        # Process the backend file
        backend_result = process_backend_file(
            backend_path, 
            auto_detect_header=auto_detect_header, 
            clean_columns=clean_columns, 
            remove_empty_rows=remove_empty_rows
        )
        
        # Clean up uploaded file after processing
        try:
            os.remove(backend_path)
            print(f"Cleaned up backend file: {backend_path}")
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up backend file {backend_path}: {cleanup_error}")
        
        # Format success/error messages
        if backend_result['success']:
            flash(f'✅ Backend file processed successfully!')
            flash(f'📁 Processed {backend_result["final_rows"]} rows from {backend_result["original_rows"]} original rows')
            flash(f'📥 ROB.xlsx ready for download!')
        else:
            flash(f'❌ Backend processing failed: {backend_result["error"]}')
        
        # Render template with backend results
        return render_template('weekly_report.html', 
                             qualified_rids=None,
                             filter_summary=None,
                             backend_result=backend_result)
        
    except Exception as e:
        print(f"Backend processing error: {e}")
        flash(f'❌ Error processing backend file: {str(e)}')
        return redirect(request.url)

def get_qualified_rids_and_remove_trending(file_path, min_search_volume, competition_level, analyze_trends=False):
    """Apply custom filters, get qualified RIDs with priority system and 1500 keyword limit"""
    try:
        print(f"Processing file: {file_path}")
        
        # Read the file
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path, engine='openpyxl')
        
        original_count = len(df)
        print(f"Original data loaded: {original_count} rows")
        
        # Validate required columns exist
        required_columns = ['AVG. Search', 'Competition', 'RID']
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            raise ValueError(f"Missing required columns: {missing_columns}")
        
        # STEP 1: Apply search volume filter first
        search_filtered_df = df[df['AVG. Search'] >= min_search_volume].copy()
        print(f"After search volume filter (>= {min_search_volume:,}): {len(search_filtered_df)} rows")
        
        # STEP 2: Apply priority-based competition filter
        print(f"Applying priority-based competition filter...")
        
        if competition_level == 'All':
            filtered_df = search_filtered_df
            print(f"No competition filter applied")
            competition_display = "All levels"
        else:
            # Priority order: Low -> Medium -> High
            priority_order = ['Low', 'Medium', 'High']
            
            # Get the selected level and all levels up to that priority
            if competition_level in priority_order:
                selected_index = priority_order.index(competition_level)
                allowed_levels = priority_order[:selected_index + 1]  # Include all levels up to selected
                
                print(f"Competition priority: {' → '.join(allowed_levels)} (up to {competition_level})")
                
                # Filter for allowed competition levels
                filtered_df = search_filtered_df[
                    search_filtered_df['Competition'].isin(allowed_levels)
                ].copy()
                
                # Sort by priority order (Low first, then Medium, then High) and search volume
                priority_map = {'Low': 1, 'Medium': 2, 'High': 3}
                filtered_df['competition_priority'] = filtered_df['Competition'].map(priority_map)
                filtered_df = filtered_df.sort_values(['competition_priority', 'AVG. Search'], ascending=[True, False])
                filtered_df = filtered_df.drop('competition_priority', axis=1)
                
                competition_display = f"Priority: {' → '.join(allowed_levels)}"
            else:
                # Fallback to exact match if level not in priority order
                filtered_df = search_filtered_df[
                    search_filtered_df['Competition'] == competition_level
                ].copy()
                competition_display = competition_level
        
        filtered_count = len(filtered_df)
        print(f"After applying priority competition filter: {filtered_count} rows")
        
        # STEP 3: Apply 1500 keyword limit BEFORE Google Trends
        max_keywords_for_trends = 1500
        
        if filtered_count > max_keywords_for_trends:
            print(f"⚠️  Too many keywords ({filtered_count}) for Google Trends analysis!")
            print(f"🔪 Limiting to top {max_keywords_for_trends} keywords (sorted by priority & search volume)")
            
            # Take top 1500 keywords (already sorted by priority and search volume)
            filtered_df = filtered_df.head(max_keywords_for_trends).copy()
            filtered_count = len(filtered_df)
            
            print(f"✅ Limited to {filtered_count} keywords for processing")
        
        # Create filter summary with priority info and limit info
        filter_summary = {
            'min_search': f"{min_search_volume:,}",
            'competition': competition_display,
            'original_count': original_count,
            'filtered_count': filtered_count,
            'trends_enabled': analyze_trends,
            'keyword_limit_applied': filtered_count == max_keywords_for_trends,
            'max_keywords_limit': max_keywords_for_trends
        }
        
        updated_file_path = None
        
        if filtered_count == 0:
            print("No records match the filter criteria")
            return [], filter_summary, updated_file_path
        
        if analyze_trends:
            # Run Google Trends analysis on filtered data (max 1500 keywords)
            print("🔥 Running Google Trends analysis on filtered keywords...")
            print(f"📊 Processing {filtered_count} keywords (within 1500 limit)")
            
            # Check if API key is configured
            if not GOOGLE_TRENDS_CONFIG.get('API_KEY') or GOOGLE_TRENDS_CONFIG['API_KEY'] == 'YOUR_API_KEY_HERE':
                print("⚠️ No API key configured - returning all filtered RIDs")
                qualified_rids = filtered_df['RID'].tolist()
                return qualified_rids, filter_summary, updated_file_path
            
            # Run actual Google Trends analysis
            keywords_data = filtered_df.to_dict('records')
            trending_data = analyze_keywords_with_google_trends(keywords_data)
            qualified_rids = [item['RID'] for item in trending_data if 'RID' in item]
            
            print(f"Google Trends analysis complete: {len(qualified_rids)} trending RIDs out of {filtered_count} filtered")
            filter_summary['trends_qualified'] = len(qualified_rids)
            filter_summary['trends_message'] = f"After Google Trends analysis: {len(qualified_rids)} out of {filtered_count} keywords are trending"
            
            # Remove trending RIDs from original dataframe
            if qualified_rids:
                print(f"🗑️ Removing {len(qualified_rids)} trending RIDs from ranking sheet...")
                
                # Create a copy of original dataframe
                df_updated = df.copy()
                
                # Remove rows where RID is in the qualified_rids list
                df_updated = df_updated[~df_updated['RID'].isin(qualified_rids)]
                
                rows_removed = len(df) - len(df_updated)
                print(f"✅ Removed {rows_removed} trending rows from ranking sheet")
                
                # Save the updated ranking sheet
                updated_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'updated_ranking_sheet.xlsx')
                
                # Use xlsxwriter for better performance
                with pd.ExcelWriter(updated_file_path, engine='xlsxwriter') as writer:
                    df_updated.to_excel(writer, index=False, sheet_name='Sheet1')
                
                print(f"💾 Updated ranking sheet saved as: {updated_file_path}")
                
                # Update filter summary with removal info
                filter_summary['rows_removed'] = rows_removed
                filter_summary['final_sheet_rows'] = len(df_updated)
                filter_summary['removal_message'] = f"Removed {rows_removed} trending rows. Updated sheet has {len(df_updated)} rows."
            else:
                print("ℹ️ No trending RIDs found - ranking sheet unchanged")
                filter_summary['removal_message'] = "No trending RIDs found - ranking sheet unchanged"
            
        else:
            # No Google Trends - return all filtered RIDs (max 1500)
            qualified_rids = filtered_df['RID'].tolist()
            print(f"Returning all filtered RIDs: {len(qualified_rids)} RIDs")
            filter_summary['trends_message'] = "Google Trends analysis disabled - showing all filtered results"
            
            if filter_summary['keyword_limit_applied']:
                filter_summary['trends_message'] += f" (limited to top {max_keywords_for_trends})"
        
        return qualified_rids, filter_summary, updated_file_path
        
    except Exception as e:
        print(f"Error in get_qualified_rids_and_remove_trending: {e}")
        raise e

def process_backend_file(file_path, auto_detect_header=True, clean_columns=True, remove_empty_rows=True):
    """Process large backend file directly to ROB format with optimization"""
    try:
        print(f"\n=== PROCESSING LARGE BACKEND FILE TO ROB FORMAT ===")
        print(f"Processing file: {file_path}")
        
        # Step 1: Read the file with optimization for large files
        try:
            # Try reading with openpyxl engine for better large file handling
            df_raw = pd.read_excel(file_path, header=None, engine='openpyxl')
        except Exception as e:
            print(f"Error with openpyxl, trying alternative: {e}")
            # Fallback to default engine
            df_raw = pd.read_excel(file_path, header=None)
        
        original_rows = df_raw.shape[0]
        print(f"Initial raw data shape: {df_raw.shape}")
        
        # Step 2: Find the actual header row if auto-detect is enabled
        if auto_detect_header:
            header_row_index = find_header_row(df_raw)
        else:
            header_row_index = 0  # Assume first row is header
        
        if header_row_index is not None:
            # Set the header
            header = df_raw.iloc[header_row_index]
            # Drop rows before the header (inclusive)
            df_data = df_raw[header_row_index + 1:].copy()
            # Assign the correct header
            df_data.columns = header
            
            # Reset index
            df_data.reset_index(drop=True, inplace=True)
            
            print(f"Data extracted with header found at index {header_row_index}. New shape: {df_data.shape}")
            
            if clean_columns:
                # Clean column names (remove leading/trailing spaces, handle duplicates)
                df_data.columns = df_data.columns.str.strip()
                
                # Handle duplicate columns
                cols = pd.Series(df_data.columns)
                for dup in cols[cols.duplicated()].unique():
                    cols[cols[cols == dup].index.values.tolist()] = [dup + '_' + str(i) if i != 0 else dup for i in range(sum(cols == dup))]
                df_data.columns = cols
                
                print("Columns cleaned.")
            
            if remove_empty_rows:
                # Drop rows that are entirely null after extraction
                initial_rows = df_data.shape[0]
                df_data.dropna(how='all', inplace=True)
                rows_dropped = initial_rows - df_data.shape[0]
                print(f"Dropped {rows_dropped} empty rows.")
            
            # Save the processed file as "ROB.xlsx" directly
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'ROB.xlsx')
            
            # Use xlsxwriter engine for better performance with large files
            with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
                df_data.to_excel(writer, index=False, sheet_name='Sheet1')
            
            print(f"ROB file saved as: {output_path}")
            
            # Display sample of processed data
            print(f"\nProcessed {len(df_data)} records successfully")
            print("Sample of processed data:")
            print(df_data.head(2).to_string())
            
            # Create summary
            summary = {
                'success': True,
                'original_rows': original_rows,
                'final_rows': len(df_data),
                'header_row': header_row_index,
                'final_columns': len(df_data.columns),
                'output_file': 'ROB.xlsx'
            }
            
            return summary
        
        else:
            print("Could not automatically detect header row.")
            return {
                'success': False,
                'error': 'Could not automatically detect header row. Please check your file format.'
            }
    
    except MemoryError:
        print("Memory error - file too large")
        return {
            'success': False,
            'error': 'File too large to process. Please try with a smaller file or contact support.'
        }
    except Exception as e:
        print(f"Error processing backend file: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def find_header_row(df):
    """Heuristic function to find the header row"""
    for index, row in df.iterrows():
        if sum(isinstance(x, str) for x in row) >= 5:
            print(f"Potential header row found at index: {index}")
            return index
    return None

@app.route('/download_backend_file')
def download_backend_file():
    """Download the processed ROB file"""
    try:
        filename = 'ROB.xlsx'
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            flash('❌ Processed file not found. Please process a backend file first.')
            return redirect(url_for('weekly_report'))
    except Exception as e:
        flash(f'❌ Error downloading file: {str(e)}')
        return redirect(url_for('weekly_report'))

@app.route('/download_updated_ranking')
def download_updated_ranking():
    """Download the updated ranking sheet (with trending rows removed)"""
    try:
        filename = 'updated_ranking_sheet.xlsx'
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name='ranking_sheet_trending_removed.xlsx')
        else:
            flash('❌ Updated ranking sheet not found. Please run Google Trends analysis first.')
            return redirect(url_for('weekly_report'))
    except Exception as e:
        flash(f'❌ Error downloading updated ranking sheet: {str(e)}')
        return redirect(url_for('weekly_report'))

# ============================================================================
# GOOGLE TRENDS FUNCTIONS - UPDATED TO USE NEW EXTRACTOR
# ============================================================================

# Updated Google Trends Config
GOOGLE_TRENDS_CONFIG = {
    'API_KEY': '68789844ac879558272d4e4d',  # ScrapingDog API key
    'INTEREST_THRESHOLD': 50,
    'DAYS_ABOVE_THRESHOLD': 2,
    'TERMS_TO_REMOVE': ['market', 'size', 'analysis', 'report', 'industry', 'global'],
    'REQUEST_DELAY': 2
}

class GoogleTrendsExtractor:
    def __init__(self, api_key):
        self.api_key = api_key
        self.base_url = "https://api.scrapingdog.com/google_trends"
    
    def get_values(self, keyword):
        """Get exact values for last 7 days"""
        params = {
            "api_key": self.api_key,
            "query": keyword,
            "geo": "US",           # Worldwide
            "tz": "330",         # Indian timezone (UTC+5:30)
            "date": "now 7-d",   # Last 7 days
            "data_type": "TIMESERIES"
        }
        try:
            response = requests.get(self.base_url, params=params)
            if response.status_code == 200:
                data = response.json()
                values = self.extract_values(data)
                return values
            else:
                print(f"    API Error: {response.status_code}")
                print(f"    Request URL: {response.url}")
                print(f"    Response Content: {response.text}")
                return []
        except Exception as e:
            print(f"    Network Error: {e}")
            return []
    
    def extract_values(self, data):
        """Extract values using standard timeline method"""
        values = []
        
        try:
            if 'interest_over_time' in data:
                timeline_data = data['interest_over_time'].get('timeline_data', [])
                
                for entry in timeline_data:
                    if isinstance(entry, dict) and 'values' in entry:
                        for val_item in entry['values']:
                            if isinstance(val_item, dict) and 'value' in val_item:
                                try:
                                    val = int(val_item['value'])
                                    if 0 <= val <= 100:
                                        values.append(val)
                                except (ValueError, TypeError):
                                    pass
        except Exception:
            pass
        
        return values
    
    def filter_keyword(self, keyword):
        """Check if keyword has 2+ values > 50 in last 7 days"""
        values = self.get_values(keyword)
        
        if values:
            count_above_50 = sum(1 for val in values if val > 50)
            print(f"    Values: {values} | Count >50: {count_above_50}")
            return count_above_50 >= 2
        
        print(f"    No values retrieved")
        return False

def analyze_keywords_with_google_trends(keywords_data):
    """Analyze keywords with Google Trends using new extractor (with 1500 keyword limit already applied)"""
    api_key = GOOGLE_TRENDS_CONFIG['API_KEY']
    extractor = GoogleTrendsExtractor(api_key)
    qualifying_keywords = []
    
    total_keywords = len(keywords_data)
    print(f"🔍 Analyzing {total_keywords} filtered keywords with Google Trends...")
    
    # Additional safety check (should already be limited to 1500)
    if total_keywords > 1500:
        print(f"⚠️ WARNING: Received {total_keywords} keywords, but max should be 1500!")
        keywords_data = keywords_data[:1500]
        print(f"🔪 Emergency limit applied: Processing only first 1500 keywords")
    
    for i, keyword_row in enumerate(keywords_data):
        try:
            original_keyword = keyword_row.get('Keywords', '')
            rid = keyword_row.get('RID', '')
            competition = keyword_row.get('Competition', '')
            search_volume = keyword_row.get('AVG. Search', 0)
            
            if not original_keyword or not rid:
                print(f"[{i+1}/{len(keywords_data)}] Skipping row with missing keyword or RID")
                continue
                
            clean_keyword = clean_keyword_for_trends(original_keyword)
            
            if not clean_keyword:
                print(f"[{i+1}/{len(keywords_data)}] Skipping empty keyword after cleaning: {original_keyword}")
                continue
            
            print(f"[{i+1}/{len(keywords_data)}] Analyzing RID {rid}: '{original_keyword}' → '{clean_keyword}' [{competition}, {search_volume:,}]")
            
            # Use new extractor to check if keyword qualifies
            is_trending = extractor.filter_keyword(clean_keyword)
            
            if is_trending:
                qualifying_keywords.append({
                    'RID': rid, 
                    'keyword': original_keyword,
                    'competition': competition,
                    'search_volume': search_volume
                })
                print(f"  ✅ TRENDING: RID {rid} - {original_keyword}")
            else:
                print(f"  ❌ Not trending: RID {rid} - {original_keyword}")
                
            # Rate limiting
            time.sleep(GOOGLE_TRENDS_CONFIG['REQUEST_DELAY'])
                
        except Exception as e:
            print(f"  ❌ Error analyzing RID {keyword_row.get('RID', 'unknown')}: {e}")
            continue
    
    print(f"🎯 Analysis complete: {len(qualifying_keywords)} out of {len(keywords_data)} keywords are trending")
    
    # Sort qualifying keywords by priority (Low first, then by search volume)
    if qualifying_keywords:
        priority_map = {'Low': 1, 'Medium': 2, 'High': 3}
        qualifying_keywords.sort(key=lambda x: (
            priority_map.get(x.get('competition', 'High'), 4),  # Priority first
            -x.get('search_volume', 0)  # Then by search volume (descending)
        ))
        print(f"📈 Qualified keywords sorted by priority and search volume")
    
    return qualifying_keywords

def clean_keyword_for_trends(keyword):
    """Clean keyword by removing problematic terms"""
    if not keyword:
        return ""
        
    cleaned = str(keyword)
    
    # Remove terms from config
    for term in GOOGLE_TRENDS_CONFIG['TERMS_TO_REMOVE']:
        cleaned = re.sub(rf'\b{re.escape(term)}\b', '', cleaned, flags=re.IGNORECASE)
    
    # Clean up extra spaces and trim
    cleaned = ' '.join(cleaned.split()).strip()
    return cleaned

# ============================================================================
# POWER AUTOMATE ROUTES
# ============================================================================

@app.route('/wait_power_automate')
def wait_power_automate():
    """Show a waiting page for Power Automate Desktop step."""
    return render_template('wait_power_automate.html')

@app.route('/api/trigger_power_automate', methods=['POST'])
def trigger_power_automate_flow():
    """Triggers a Power Automate Desktop flow"""
    pad_exe_path = r"C:\Program Files (x86)\Power Automate Desktop\PAD.Console.Host.exe"
    flow_name = "Paid PR - Files Downloader"
    
    if not os.path.exists(pad_exe_path):
        print("Power Automate Desktop executable not found!")
        return jsonify({'status': 'error', 'message': 'PAD executable not found'})
    
    command = f'"{pad_exe_path}" -flow "{flow_name}"'
    
    try:
        result = subprocess.run(command, shell=True, check=True, text=True, capture_output=True)
        print(f"Flow triggered successfully. Output: {result.stdout}")

        time.sleep(5)
        
        flow_button_coordinates = (463, 395)
        print(f"Clicking at {flow_button_coordinates}")
        pyautogui.click(flow_button_coordinates)
        print("Flow triggered successfully.")

    except subprocess.CalledProcessError as e:
        print(f"Error triggering flow: {e.stderr}")
        return jsonify({'status': 'error', 'message': f'Flow error: {e.stderr}'})
    
    return jsonify({'status': 'success', 'message': 'Power Automate process completed.'})

# ============================================================================
#  Custom APPLICATION RUNNER
# ============================================================================
@app.route('/custom_index.html')
def custom_index():
    """Render custom index page for application"""
    return render_template('custom_index.html')



@app.route('/custom_weekly_report', methods=['GET', 'POST'])
def custom_weekly_report():
    if request.method == 'POST':
        try:
            print("Custom Weekly RID Analysis POST request received!")
            
            # Get filter parameters from form
            min_search_volume = int(request.form.get('min_search_volume', 5000))
            competition_level = request.form.get('competition_level', 'Low')
            analyze_trends = request.form.get('analyze_trends') == 'on'
            
            print(f"Custom Weekly Filters: Search >= {min_search_volume}, Competition = {competition_level}")
            print(f"Google Trends: {'Enabled' if analyze_trends else 'Disabled'}")
            
            # Validate form inputs
            if not min_search_volume or min_search_volume < 0:
                flash('❌ Please enter a valid minimum search volume!')
                return redirect(request.url)
                
            if not competition_level:
                flash('❌ Please select a competition level!')
                return redirect(request.url)
            
            # Handle RANKING SHEET upload only
            ranking_file = request.files.get('ranking_file')
            if not ranking_file or ranking_file.filename == '':
                flash('❌ Please select a ranking Excel file!')
                return redirect(request.url)

            if not allowed_file(ranking_file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed for ranking sheet!')
                return redirect(request.url)

            # Save uploaded file
            ranking_filename = secure_filename(ranking_file.filename)
            ranking_path = os.path.join(app.config['UPLOAD_FOLDER'], ranking_filename)
            ranking_file.save(ranking_path)
            print(f"Ranking file saved: {ranking_path}")
            
            # Process the ranking file and get qualified RIDs
            result_summary = process_custom_weekly_ranking_file(
                ranking_path, min_search_volume, competition_level, analyze_trends
            )
            
            # Format success/warning messages based on results
            if result_summary['success']:
                flash(f'✅ Success! Found {result_summary["qualified_rids_count"]} qualified RIDs')
                flash(f'📁 Custom weekly RID analysis completed!')
                print(f"Custom weekly analysis completed: {result_summary}")
            else:
                flash(f'❌ Error: {result_summary.get("error", "Unknown error")}')
                result_summary = None
            
            # Clean up uploaded file after processing
            try:
                os.remove(ranking_path)
                print(f"Cleaned up uploaded file")
            except Exception as cleanup_error:
                print(f"Warning: Could not clean up file: {cleanup_error}")
            
            # Render template with results
            return render_template('custom_weekly_report.html', 
                                  qualified_rids=result_summary.get('qualified_rids', []) if result_summary else [],
                                  filter_summary=result_summary.get('filter_summary', {}) if result_summary else {},
                                  custom_weekly_result=result_summary)
            
        except ValueError as ve:
            print(f"Value Error: {ve}")
            flash('❌ Invalid input values. Please check your filters.')
            return redirect(request.url)
        except Exception as e:
            print(f"Error: {e}")
            flash(f'❌ Error processing file: {str(e)}')
            return redirect(request.url)
    
    # GET request - show custom weekly form
    return render_template('custom_weekly_report.html')

def process_custom_weekly_ranking_file(ranking_path, min_search_volume, competition_level, analyze_trends):
    """Process ranking file for custom weekly analysis (no ROB matching)"""
    try:
        print(f"\n=== PROCESSING CUSTOM WEEKLY RANKING FILE ===")
        print(f"Ranking file: {ranking_path}")
        
        # Process ranking sheet to get qualified RIDs
        qualified_rids, filter_summary, updated_ranking_path = get_qualified_rids_and_remove_trending(
            ranking_path, min_search_volume, competition_level, analyze_trends
        )
        
        if not qualified_rids:
            return {
                'success': False,
                'error': 'No qualified RIDs found in ranking sheet with your filter criteria'
            }
        
        print(f"✅ Found {len(qualified_rids)} qualified RIDs for custom weekly analysis")
        
        # Read the original ranking file to get RID and Keywords data
        if ranking_path.endswith('.csv'):
            df_original = pd.read_csv(ranking_path)
        else:
            df_original = pd.read_excel(ranking_path, engine='openpyxl')
        
        print(f"Original file columns: {list(df_original.columns)}")
        
        # Convert qualified_rids to same type as RID column for matching
        df_original['RID'] = df_original['RID'].astype(str).str.strip()
        qualified_rids_str = [str(rid).strip() for rid in qualified_rids]
        
        # Filter rows that match qualified RIDs and select only RID and Keywords columns
        qualified_data = df_original[df_original['RID'].isin(qualified_rids_str)][['RID', 'Keywords']].copy()
        
        print(f"Qualified data shape: {qualified_data.shape}")
        print(f"Sample data:\n{qualified_data.head()}")
        
        # Save to Desktop/RPA folder as Excel file
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        output_path = os.path.join(rpa_folder, 'Custom_weekly_ROB.xlsx')
        
        # Export to Excel - CORRECTED VERSION
        qualified_data.to_excel(output_path, index=False, engine='xlsxwriter')
        
        print(f"✅ Custom weekly ROB Excel saved to: {output_path}")
        print(f"Excel contains {len(qualified_data)} rows with RID and Keywords")
        
        # Verify file was created
        if os.path.exists(output_path):
            print(f"✅ File verification: Excel file exists at {output_path}")
        else:
            print(f"❌ File verification: Excel file NOT found at {output_path}")
        
        # Create comprehensive summary
        summary = {
            'success': True,
            'qualified_rids': qualified_rids,
            'qualified_rids_count': len(qualified_rids),
            'output_path': output_path,
            'filter_summary': filter_summary,
            'analysis_type': 'custom_weekly',
            'excel_rows': len(qualified_data)
        }
        
        return summary
        
    except Exception as e:
        print(f"Error in custom weekly ranking file processing: {e}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'error': str(e)
        }
    

@app.route('/custom_choice')
def custom_choice():
    """Custom choice page for CMI/WMR selection"""
    return render_template('custom_choice.html')



@app.route('/custom_cmi_cta', methods=['GET', 'POST'])
def custom_cmi_cta():
    if request.method == 'POST':
        try:
            # Get form data (only extract_count now since user info fields are removed)
            extract_count = int(request.form.get('extract_count', 200))

            # Validate required fields (only extract_count and file now)
            if not extract_count or extract_count < 1:
                flash('❌ Please enter a valid number of keywords to extract!')
                return redirect(request.url)

            file = request.files.get('file')
            if not file or file.filename == '':
                flash('❌ Keywords file is required!')
                return redirect(request.url)

            if not allowed_file(file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed!')
                return redirect(request.url)

            # Use secure_filename to avoid path issues
            filename = secure_filename(file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(input_path)

            # Process the file for CMI CTA generation
            result = process_cmi_cta_file(input_path, extract_count)
            
            if result['success']:
                flash(f'✅ Successfully processed {result["extracted_count"]} keywords!')
                flash(f'📁 File saved: {result["filename"]}')
                flash('🤖 CMI automation started!')
                
                # Start CMI automation in background
                threading.Thread(target=run_cmi_automation).start()
            else:
                flash(f'❌ Error: {result["error"]}')

            return render_template('custom_cmi_cta.html')

        except ValueError as ve:
            flash('❌ Please enter a valid number for keywords to extract!')
            return redirect(request.url)
        except Exception as e:
            flash(f'❌ Error processing file: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_cmi_cta.html')

def process_cmi_cta_file(file_path, extract_count):
    """Process Custom Weekly ROB file for CMI CTA generation"""
    try:
        # Read the file
        if file_path.endswith('.csv'):
            df_original = pd.read_csv(file_path)
        else:
            df_original = pd.read_excel(file_path, engine='openpyxl')

        total_rows = len(df_original)
        
        if total_rows < extract_count:
            extract_count = total_rows

        # Step 1: Extract top N rows for CTA generation
        extracted_rows = df_original.head(extract_count).copy()
        
        # Step 2: Get remaining rows (original minus extracted)
        remaining_rows = df_original.iloc[extract_count:].copy()
        
        # Create filename with current date for extracted data
        today = datetime.today()
        extracted_filename = f"custom_reports_cmi_{today.year}_{today.month:02d}_{today.day:02d}.xlsx"
        
        # Save extracted data to Desktop/RPA folder
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        # Save extracted keywords file
        extracted_output_path = os.path.join(rpa_folder, extracted_filename)
        extracted_rows.to_excel(extracted_output_path, index=False)
        
        # Step 3: Update the original Custom_weekly_ROB.xlsx file with remaining data
        custom_weekly_rob_path = os.path.join(rpa_folder, 'Custom_weekly_ROB.xlsx')
        
        if os.path.exists(custom_weekly_rob_path):
            # Update the original file with remaining rows
            remaining_rows.to_excel(custom_weekly_rob_path, index=False)
            print(f"✅ Updated Custom_weekly_ROB.xlsx - Removed {extract_count} extracted keywords")
            print(f"✅ Custom_weekly_ROB.xlsx now contains {len(remaining_rows)} remaining keywords")
        else:
            print(f"⚠️ Warning: Custom_weekly_ROB.xlsx not found at {custom_weekly_rob_path}")
        
        return {
            'success': True,
            'extracted_count': extract_count,
            'remaining_count': len(remaining_rows),
            'total_count': total_rows,
            'filename': extracted_filename,
            'output_path': extracted_output_path,
            'updated_original': os.path.exists(custom_weekly_rob_path)
        }
        
    except Exception as e:
        print(f"Error in process_cmi_cta_file: {e}")
        return {
            'success': False,
            'error': str(e)
        }


def run_cmi_automation():
    """Run CMI automation using Selenium"""
    try:
        print("Starting CMI automation...")
        
        from webdriver_manager.chrome import ChromeDriverManager
        from selenium import webdriver
        from selenium.webdriver.chrome.service import Service
        from selenium.webdriver.common.by import By
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.webdriver.support.ui import Select
        import time
        
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        #options.add_argument("--headless")  # Run headlessly for deployment
        options.add_argument("--start-maximized")
        
        # Initialize WebDriver
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.coherentmarketinsights.com/cmisitmanup/index.php')
        
        username_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH,'//*[@id="adlogin"]/div[1]/input'))
        )
        username_input.send_keys('Auto_Ops_Team')
        
        password_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH,'//*[@id="adlogin"]/div[2]/input'))
        )
        password_input.send_keys('kDp7%8^03Ib')
        
        signup_click = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH,'//*[@id="adlogin"]/div[3]/div/button'))
        )
        signup_click.click()
        
        custom_insights_click = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH,'/html/body/div/aside/section/ul/li[3]/a/span[1]'))
        )
        custom_insights_click.click()
        
        print("CMI automation completed successfully!")
        
        # Keep browser open for now - you can modify this behavior
        time.sleep(10)
        # driver.quit()  # Uncomment to close browser automatically
        
    except Exception as e:
        print(f"CMI automation error: {e}")


@app.route('/custom_wmr_cta', methods=['GET', 'POST'])
def custom_wmr_cta():
    if request.method == 'POST':
        try:
            # Get form data (only extract_count now since user info fields are removed)
            extract_count = int(request.form.get('extract_count', 200))

            # Validate required fields (only extract_count and file now)
            if not extract_count or extract_count < 1:
                flash('❌ Please enter a valid number of keywords to extract!')
                return redirect(request.url)

            file = request.files.get('file')
            if not file or file.filename == '':
                flash('❌ Keywords file is required!')
                return redirect(request.url)

            if not allowed_file(file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed!')
                return redirect(request.url)

            # Use secure_filename to avoid path issues
            filename = secure_filename(file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(input_path)

            # Process the file for WMR CTA generation
            result = process_wmr_cta_file(input_path, extract_count)
            
            if result['success']:
                flash(f'✅ Successfully processed {result["extracted_count"]} keywords!')
                flash(f'📁 File saved: {result["filename"]}')
                flash('🤖 WMR automation started!')
                
                # Start WMR automation in background
                threading.Thread(target=run_wmr_automation).start()
            else:
                flash(f'❌ Error: {result["error"]}')

            return render_template('custom_wmr_cta.html')

        except ValueError as ve:
            flash('❌ Please enter a valid number for keywords to extract!')
            return redirect(request.url)
        except Exception as e:
            flash(f'❌ Error processing file: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_wmr_cta.html')


def process_wmr_cta_file(file_path, extract_count):
    """Process Custom Weekly ROB file for WMR CTA generation"""
    try:
        # Read the file
        if file_path.endswith('.csv'):
            df_original = pd.read_csv(file_path)
        else:
            df_original = pd.read_excel(file_path, engine='openpyxl')

        total_rows = len(df_original)
        
        if total_rows < extract_count:
            extract_count = total_rows

        # Step 1: Extract top N rows for WMR CTA generation
        extracted_rows = df_original.head(extract_count).copy()
        
        # Step 2: Get remaining rows (original minus extracted)
        remaining_rows = df_original.iloc[extract_count:].copy()
        
        # Create filename with current date for extracted data
        today = datetime.today()
        extracted_filename = f"custom_reports_wmr_{today.year}_{today.month:02d}_{today.day:02d}.xlsx"
        
        # Save extracted data to Desktop/RPA folder
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        # Save extracted keywords file
        extracted_output_path = os.path.join(rpa_folder, extracted_filename)
        extracted_rows.to_excel(extracted_output_path, index=False)
        
        # Step 3: Update the original Custom_weekly_ROB.xlsx file with remaining data
        custom_weekly_rob_path = os.path.join(rpa_folder, 'Custom_weekly_ROB.xlsx')
        
        if os.path.exists(custom_weekly_rob_path):
            # Update the original file with remaining rows
            remaining_rows.to_excel(custom_weekly_rob_path, index=False)
            print(f"✅ Updated Custom_weekly_ROB.xlsx - Removed {extract_count} extracted keywords for WMR")
            print(f"✅ Custom_weekly_ROB.xlsx now contains {len(remaining_rows)} remaining keywords")
        else:
            print(f"⚠️ Warning: Custom_weekly_ROB.xlsx not found at {custom_weekly_rob_path}")
        
        return {
            'success': True,
            'extracted_count': extract_count,
            'remaining_count': len(remaining_rows),
            'total_count': total_rows,
            'filename': extracted_filename,
            'output_path': extracted_output_path,
            'updated_original': os.path.exists(custom_weekly_rob_path)
        }
        
    except Exception as e:
        print(f"Error in process_wmr_cta_file: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def run_wmr_automation():
    """Run WMR automation using Selenium with your provided code"""
    try:
        print("Starting WMR automation...")
        
        from webdriver_manager.chrome import ChromeDriverManager
        from selenium import webdriver
        from selenium.webdriver.chrome.service import Service
        from selenium.webdriver.common.by import By
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.webdriver.support.ui import Select
        import time

        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        #options.add_argument("--headless")  # Run headlessly for deployment
        options.add_argument("--start-maximized")

        # Initialize WebDriver
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.worldwidemarketreports.com/imanagereports')
                
        username_input = WebDriverWait(driver, 10).until(
             EC.presence_of_element_located((By.XPATH,'//*[@id="adlogin"]/div[1]/input'))
            )
        username_input.send_keys('Auto_Ops_Team')
                
        password_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH,'//*[@id="adlogin"]/div[2]/input'))
            )
        password_input.send_keys('M9b@0j9Y28O')
                
        login_click = WebDriverWait(driver, 10).until(
          EC.element_to_be_clickable((By.XPATH,'//*[@id="adlogin"]/div[3]/div/button'))
            )
        login_click.click()
                
        custom_insights_click = WebDriverWait(driver, 10).until(
           EC.element_to_be_clickable((By.XPATH,'/html/body/div/aside/section/ul/li[3]/a/span[1]'))
            )
        custom_insights_click.click()
        
        print("WMR automation completed successfully!")
        
        # Keep browser open for now - you can modify this behavior
        time.sleep(10)
        # driver.quit()  # Uncomment to close browser automatically
        
    except Exception as e:
        print(f"WMR automation error: {e}")


@app.route('/custom_content_generation_choice')
def custom_content_generation_choice():
    """Custom content generation choice page"""
    return render_template('custom_content_generation_choice.html')

# Add these imports at the top if not already present
import openai
from docx import Document
import re

# Configure OpenAI (add your API key)
OPENAI_API_KEY = "sk-osX5I2lupISEx3asBA8gO71GgcmPp7mIcUWoczALHVT3BlbkFJF72ozu5mUnkKmLgiSepX7n5Fd-UmvHC5g1JhBA33YA"  # Replace with your actual API key

@app.route('/custom_ai_content', methods=['GET', 'POST'])
def custom_ai_content():
    if request.method == 'POST':
        try:
            # Handle file upload
            cta_file = request.files.get('cta_file')
            if not cta_file or cta_file.filename == '':
                flash('❌ CTA excel file is required!')
                return redirect(request.url)

            if not allowed_file(cta_file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) are allowed!')
                return redirect(request.url)

            # Save uploaded file
            filename = secure_filename(cta_file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            cta_file.save(input_path)

            # Process the file for AI content generation
            result = process_ai_content_generation(input_path)
            
            if result['success']:
                flash(f'✅ Successfully generated {result["articles_created"]} AI articles!')
                flash(f'📁 Articles saved to Desktop/RPA folder')
            else:
                flash(f'❌ Error: {result["error"]}')

            # Clean up uploaded file
            try:
                os.remove(input_path)
            except:
                pass

            return render_template('custom_ai_content.html')

        except Exception as e:
            flash(f'❌ Error processing file: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_ai_content.html')

def clean_title(title):
    """Remove 'Market' and related words from title"""
    # Remove common market-related terms
    market_terms = [
        r'\bmarket\b', r'\bMarket\b', r'\bMARKET\b',
        r'\bmarket size\b', r'\bMarket Size\b',
        r'\bmarket analysis\b', r'\bMarket Analysis\b',
        r'\bmarket research\b', r'\bMarket Research\b',
        r'\bmarket report\b', r'\bMarket Report\b',
        r'\bmarket study\b', r'\bMarket Study\b'
    ]
    
    cleaned_title = title
    for term in market_terms:
        cleaned_title = re.sub(term, '', cleaned_title, flags=re.IGNORECASE)
    
    # Clean up extra spaces and punctuation
    cleaned_title = re.sub(r'\s+', ' ', cleaned_title).strip()
    cleaned_title = re.sub(r'^[-\s]+|[-\s]+$', '', cleaned_title)
    
    return cleaned_title

def generate_article_with_openai(clean_title, promo_link, sample_link):
    """Generate article using OpenAI API"""
    try:
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        
        prompt = f"""

We are coherent market insights and going to publish this data on third part PR website, thus do not mention any competitors name (market intelligence firm, market research, consulting and data and analytics companies), and take ownership for market share, size, growth rate as this research done by our own research team however any real instances and stats you can use reference such as company or organisation or others as mentioned in detail explanation. Write the blog covering the below pointers: Start content with market size and overview paragraph (do not add heading and do not use bullet for this section), ➔ add first CTA link here, ➤Actionable Insights, ➤Market Segment and Regional coverage, ➔ add Second CTA link here, ➤Key players, ➤Growth factors, ➤ Market trends, ➤Key takeaways, and ❓ Frequently Asked Questions, All these pointers should act as a heading to respective paragraphs, do not miss any bullet foe above given. Consider the following points while generating content: Flow of information, all given secondary keywords must be covered, and there should be a heading for each paragraph or bullet pointers. Try to cover content in in bullet pointers whenever possible each paragraph should be short, about 100 to 120 words. Our readers are already experts in the field, so always try to generate content that provides unique insights and value addition for experts. Thus, while generating each piece of content, it should be data-backed with actual instances and stats from recent years 2025 and 2024, covering maximum stats that increase the authenticity and are enough to support or make decisions based upon reading this blog. Do not add generic content that is already known to readers, such as definitions, advantages, disadvantages, or other generic content. Flow of information should be as below. Start the blog with writing 5 to 7 sentence ling paragraph start content with talking about Industry in 2 to 3 sentences and should be matches with heading of the blog. followed by 2 to 4 sentence on market size and CAGR formatted as The Global (Insert Market Name) Market size is estimated to be valued at USD (Identify and Insert Market Value for 2025) in 2025 (market size) and is expected to reach USD (Identify and Insert Market Value for 2032) by 2032 (market forecast), exhibiting a compound annual growth rate (CAGR) of (Identify and Insert CAGR)% from 2025 to 2032. Do not add heading and do not use bullet for this section. (Make sure all values such as market size, CAGR, key players should be carefully identified with research approach).  Then First CTA link. Then Actionable Insights: In Actionable Insights, include essential quantitative indicators that lead to market size. For example, supply-side indicators (Production Capacity, Pricing, and Exports), demand-side indicators (Pricing, Imports, Various Use Cases across industries), micro-indicators, nano-size indicators—focus more on the quantitative aspects. Add stats or actual instance examples from the recent year to support the given heading. Next part is Market segment and regional Coverage where enlist the all subsegment under each segment categories and fragment region into given format. Comprehensive Segmentation and Classification of the Report: » By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. » By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. » By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. Regional and Country Analysis: » North America: U.S. and Canada » Latin America: Brazil, Argentina, Mexico, and Rest of Latin America » Europe: Germany, U.K., Spain, France, Italy, Benelux, Denmark, Norway, Sweden, Russia, and Rest of Europe » Asia Pacific: China, Taiwan, India, Japan, South Korea, Indonesia, Malaysia, Philippines, Singapore, Australia, and Rest of Asia Pacific » Middle East & Africa: Bahrain, Kuwait, Oman, Qatar, Saudi Arabia, United Arab Emirates, Israel, South Africa, North Africa, Central Africa, and Rest of MEA. Then Second CTA link. Then key Players: List 12 to 20 highly relevant key players for the given market. Furthermore, add 2-3 statements on competitive strategies adopted by a few key players, mentioning actual strategies and entities involved along with the actual outcome. Growth Factors: Growth factor heading and short description with supporting stats or examples from the recent year in the content. Then Market Trends: Market Trend heading and short description with supporting stats or examples from the recent year in the content. Then Key takeaways: Add total 6 key takeaways covering one line of description on each point.  Segment Covers (2-3 segments and sub-segments under each segment): Enlist 2-3 segments and all subsegments that fall under each segment for the given market. Mention the fastest-growing sub-segment and dominating sub-segment under each segment along with an instance from a recent year 2025 and 2024 that will support the given statement. In continuation enlist 2-3 insightful regions: Mention the Dominating region, fastest growing region, and add actual instances to support both the dominating and fastest-growing regions. Dominating Region: List one dominating region for the given market from North America, Latin America, Europe, Asia Pacific, Middle East, and Africa. Fastest Growing Region: List one of the fastest-growing regions for the given market from the above-mentioned regions. Make sure to not to mention our organization name or relevant terms anywhere in the output such as coherent market insights or our analyst team or our research team. Given Market Name and Data:  


.From an SEO perspective, we need to cover all given keywords from the list below. However, they should appear naturally so that the content flow looks natural for the reader. Keyword List: market share, market size, market research, market insights, market trends, market opportunities, market challenges, market growth, market forecast, market companies, market players, market analysis, market drivers, market restraints, market scope, market dynamics, market segments, market report, market growth strategies, market revenue, industry size, industry share, industry trends, and business growth, furthermore - Market size and market report, market revenue, market share, trends keywords are mandatory to be added twice in content. In addition to the above requirement, in 5 places, add the actual market name along with the above keywords so that long-tail keywords will be generated. These long-tail keywords are market name + size, market name + report, market name + revenue, market name + share, market name + trends. Make sure all given keywords are naturally fit, do not try to infuse forcefully, flow of information should be natural and meaningful, furthermore make sure spelling and structure of sentences from generated output are grammatically correct. Furthermore, based on the market name, create a set of Frequently Asked Questions that are highly relevant and customized to the specific market. The sample Frequently Asked Questions below are for understanding purposes only. For the given market, questions can be completely replaced. However, please tailor the actual questions to the market name and the insights provided in the report: 1. Who are the dominant players in the (Market Name) market? 2. What will be the size of the (Market Name) market in the coming years? 3. Which end users industry has the largest growth opportunity? 4. How will market development trends evolve over the next five years? 5. What is the nature of the competitive landscape and challenges in the (Market Name) market? 6. What go-to-market strategies are commonly adopted in the (Market Name) market? Make sure to answer to all FAQs. In the case of country-level markets, please exclude the word 'Global' and key takeaways where other regions are mentioned. Make sure to add catchy bullet in generated output. I have shared the reference bullet with you. Make sure to add this bullet. For heading use these bullet- ➤Actionable Insights, ➤Market Segment and Regional Coverage, ➔ Inserted Second CTA link, ➤Key Players, ➤Growth factors, ➤ Market Trends, ➤Key Takeaways, and ❓ Frequently Asked Questions. Make sure do not miss any Bullet including CTA bullet which is ➔. For subpointers under main headings use bullets which is in reference as provided- Actionable Insights ●, Market Segment and Regional Coverage●, Key players●, Growth Factors●,  Market Trends●, Key Takeaways●. Make sure to use these bullets for given subpointers. Make sure bullet appear properly and there should be single bullet for one points, ensure there no two adjacent bullet immediate to next to each other.
"""
        
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": "You are a professional market research writer specializing in industry analysis articles."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        print(f"OpenAI error: {e}")
        return f"""
{clean_title} - Market Analysis

Market Overview:
The {clean_title} industry represents a significant segment in today's market landscape. This sector has shown remarkable growth potential and continues to attract attention from investors and industry analysts.

Key Players:
Leading companies in this space are driving innovation and market expansion. For detailed market insights and analysis, industry professionals can access comprehensive research at {promo_link}.

Growth Factors:
Several factors contribute to the growth of this market, including technological advancements, increasing demand, and strategic market positioning.

Market Trends:
Current trends indicate strong market momentum with various opportunities for stakeholders. Additional market data and samples are available through {sample_link}.

Key Takeaways:
The {clean_title} sector presents substantial opportunities for growth and investment in the coming years.
"""

def save_article_as_doc(article_content, clean_title):
    """Save article as .doc file"""
    try:
        # Create Word document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(f"{clean_title} - Market Analysis", level=1)
        
        # Add content paragraphs
        paragraphs = article_content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Create filename
        today = datetime.today()
        safe_title = re.sub(r'[^\w\s-]', '', clean_title.lower())
        safe_title = re.sub(r'[-\s]+', '_', safe_title)
        filename = f"{safe_title}_cmi_{today.year}_{today.month:02d}_{today.day:02d}.doc"
        
        # Save to Desktop/RPA
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        file_path = os.path.join(rpa_folder, filename)
        doc.save(file_path)
        
        return True, filename
        
    except Exception as e:
        print(f"Error saving document: {e}")
        return False, str(e)

def process_ai_content_generation(file_path):
    """Process CTA excel file and generate AI articles"""
    try:
        print(f"\n=== PROCESSING AI CONTENT GENERATION ===")
        print(f"File: {file_path}")
        
        # Read the excel file
        df = pd.read_excel(file_path, engine='openpyxl')
        
        print(f"Excel columns: {list(df.columns)}")
        print(f"Found {len(df)} rows to process")
        
        # Verify required columns exist
        required_columns = ['Title', 'PromoBuy', 'SampleLink']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            return {
                'success': False,
                'error': f'Missing required columns: {missing_columns}'
            }
        
        articles_created = 0
        
        # Process each row
        for index, row in df.iterrows():
            try:
                original_title = str(row['Title'])
                promo_link = str(row['PromoBuy'])
                sample_link = str(row['SampleLink'])
                
                print(f"\n[{index+1}/{len(df)}] Processing: {original_title}")
                
                # Clean the title - FIXED: Use different variable name
                cleaned_title = clean_title(original_title)  # ✅ FIXED!
                print(f"Cleaned title: {cleaned_title}")
                
                # Generate article using OpenAI
                print("Generating article with OpenAI...")
                article_content = generate_article_with_openai(cleaned_title, promo_link, sample_link)
                
                # Save as .doc file
                success, filename = save_article_as_doc(article_content, cleaned_title)
                
                if success:
                    print(f"✅ Article saved: {filename}")
                    articles_created += 1
                else:
                    print(f"❌ Failed to save article: {filename}")
                
                # Small delay to avoid API rate limits
                time.sleep(1)
                
            except Exception as e:
                print(f"❌ Error processing row {index+1}: {e}")
                continue
        
        return {
            'success': True,
            'articles_created': articles_created,
            'total_rows': len(df)
        }
        
    except Exception as e:
        print(f"Error in AI content generation: {e}")
        return {
            'success': False,
            'error': str(e)
        }

if __name__ == '__main__':
    import webbrowser
    webbrowser.open('http://127.0.0.1:5000/')
    app.run(debug=True, host='0.0.0.0', port=5000)